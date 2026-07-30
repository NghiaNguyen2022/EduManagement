from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_PATH = ROOT / "scripts" / "build-user-guide-docx.py"
SPEC = spec_from_file_location("guide_base", BASE_PATH)
base = module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(base)

OUT = ROOT / "docs" / "HUONG_DAN_VAN_HANH_TUAN_TU_QLTRUONGHOC.docx"


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QLTRUONGHOC")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(base.BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HƯỚNG DẪN VẬN HÀNH\nTHEO TRÌNH TỰ CÔNG VIỆC")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Từ khởi tạo tài khoản đến tuyển sinh, đào tạo, tài chính và Portal")
    doc.add_paragraph()
    base.add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Tác giả", "Nhóm phát triển QLTruongHoc"),
            ("Phiên bản", "3.2"),
            ("Ngày cập nhật", "30/07/2026"),
            ("Đối tượng", "Quản trị, quản lý đơn vị, tuyển sinh, học vụ, giáo viên, kế toán và phụ huynh"),
            ("Phạm vi demo", "Trường Mầm non Hoa Nắng và Trung tâm Ngoại ngữ Quận 8"),
        ],
        [1.55, 4.95],
    )
    doc.add_page_break()


def procedure(
    doc,
    number,
    title,
    actor,
    prerequisites,
    purpose,
    steps,
    image,
    expected,
    tests,
    note=None,
):
    doc.add_heading(f"Bước {number}. {title}", level=2)
    doc.add_paragraph(purpose)
    base.add_callout(
        doc,
        "Ai thực hiện:",
        f"{actor}  Dữ liệu cần có: {prerequisites}",
    )
    doc.add_heading("Thao tác chi tiết", level=3)
    base.add_steps(doc, steps)
    if note:
        base.add_callout(doc, "Lưu ý nghiệp vụ:", note, base.AMBER)
    base.add_figure(doc, image, f"Hình {number}. {title}")
    doc.add_heading("Kết quả cần kiểm tra", level=3)
    for item in expected:
        doc.add_paragraph("• " + item, style="Guide Bullet")
    doc.add_heading("Test case", level=3)
    base.add_test_cases(doc, tests)


def module_guide(doc, number, title, owner, feature_rows, preschool_case, center_case, tests):
    doc.add_heading(f"4.{number}. {title}", level=2)
    base.add_callout(doc, "Vai trò phụ trách:", owner, base.GREEN)
    base.add_table(
        doc,
        ["Chức năng", "Trường mầm non", "Trung tâm ngoại ngữ"],
        feature_rows,
        [1.45, 2.55, 2.5],
    )
    doc.add_heading("Business case A — Trường mầm non", level=3)
    base.add_steps(doc, preschool_case)
    doc.add_heading("Business case B — Trung tâm ngoại ngữ", level=3)
    base.add_steps(doc, center_case)
    doc.add_heading("Test case và tiêu chí nghiệm thu", level=3)
    base.add_test_cases(doc, tests)


def build():
    doc = Document()
    base.configure_document(doc)
    cover(doc)

    doc.add_heading("Kiểm soát tài liệu", level=1)
    base.add_table(
        doc,
        ["Phiên bản", "Ngày", "Nội dung cập nhật"],
        [
            ("3.2", "30/07/2026", "Mở rộng thành cẩm nang chi tiết theo nhóm chức năng, vai trò và hai loại đơn vị; bổ sung danh mục chức năng nhỏ."),
            ("3.1", "30/07/2026", "Bổ sung quy tắc tài khoản giáo viên theo họ tên, form nhập học/thu tiền, phiếu in và bản đồ chức năng theo vai trò, loại đơn vị."),
            ("3.0", "29/07/2026", "Sắp xếp lại toàn bộ hướng dẫn theo trình tự công việc và đặt ảnh ngay tại từng bước."),
        ],
        [1.0, 1.2, 4.3],
    )
    base.add_callout(
        doc,
        "Cách sử dụng:",
        "Thực hiện tuần tự từ Bước 1. Chỉ chuyển sang bước sau khi các test case của bước hiện tại đạt.",
        base.GREEN,
    )
    doc.add_heading("Mục lục", level=1)
    base.add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    doc.add_heading("1. Sơ đồ trình tự triển khai", level=1)
    base.add_table(
        doc,
        ["Giai đoạn", "Vai trò chính", "Kết quả đầu ra"],
        [
            ("A. Chuẩn bị tổ chức", "Quản trị/Quản lý", "Tài khoản nhân sự và hồ sơ giáo viên"),
            ("B. Tuyển sinh", "Tuyển sinh/Tư vấn", "Lead, lịch sử chăm sóc, học sinh và phụ huynh"),
            ("C. Thiết lập đào tạo", "Học vụ", "Chương trình, lớp, phân công, xếp lớp và lịch"),
            ("D. Vận hành lớp", "Giáo viên/Học vụ", "Buổi học, điểm danh, báo giảng, xin phép và trao đổi"),
            ("E. Tài chính", "Kế toán/Quản lý", "Kỳ thu, công nợ, phiếu thu, điều chỉnh và chi phí"),
            ("F. Kết quả và Portal", "Giáo viên/Học vụ/Phụ huynh", "Đánh giá, chứng chỉ, thông báo và báo cáo"),
        ],
        [1.55, 1.65, 3.3],
    )
    base.add_callout(
        doc,
        "Quy tắc dữ liệu lặp:",
        "Tài liệu mô tả và chụp đầy đủ một lần tạo mẫu. Khi cần 10 học sinh, 5 giáo viên hoặc nhiều khoản thu, lặp lại đúng quy trình và thay dữ liệu đầu vào.",
    )

    doc.add_heading("2. Bản đồ chức năng theo vai trò", level=1)
    doc.add_paragraph(
        "Phần này giúp người dùng xác định nhanh công việc của mình. Khi triển khai lần đầu, "
        "vẫn thực hiện theo trình tự tại Phần 3 để dữ liệu đầu vào của bước sau luôn đầy đủ."
    )
    base.add_table(
        doc,
        ["Nhóm vai trò", "Công việc chính", "Kết quả bàn giao"],
        [
            ("Quản trị/Quản lý đơn vị", "Khởi tạo nhân sự, kiểm soát đơn vị, duyệt điều chỉnh và theo dõi báo cáo.", "Đủ tài khoản, đúng quyền, dữ liệu tách biệt theo đơn vị."),
            ("Tuyển sinh/Tư vấn", "Tiếp nhận khách hàng tiềm năng, liên hệ, đặt hẹn, xác nhận đăng ký.", "Hồ sơ học sinh/học viên và phụ huynh được tạo đúng nguồn."),
            ("Học vụ", "Tạo chương trình, lớp, giáo viên, xếp lớp, thời khóa biểu và xử lý xin phép.", "Lớp có giáo viên, học viên và buổi học thực tế."),
            ("Giáo viên", "Điểm danh, báo giảng, nhận xét, kết quả học tập và trao đổi phụ huynh.", "Dữ liệu chuyên môn gắn đúng lớp được phân công."),
            ("Kế toán", "Tạo khoản thu, kỳ thu, công nợ, thu tiền, phiếu thu và chi phí.", "Chứng từ, công nợ và báo cáo khớp nhau."),
            ("Phụ huynh", "Theo dõi con, lịch học, học phí, kết quả, xin phép, trao đổi và thông báo.", "Chỉ nhìn thấy đúng các con và đúng dữ liệu từng đơn vị."),
        ],
        [1.45, 3.35, 1.7],
    )

    doc.add_heading("2.1. Triển khai song song theo loại đơn vị", level=2)
    base.add_table(
        doc,
        ["Nhóm dữ liệu", "Trường mầm non", "Trung tâm ngoại ngữ"],
        [
            ("Chương trình", "Theo độ tuổi, lĩnh vực phát triển và kế hoạch năm học.", "Theo ngoại ngữ, cấp độ, khóa học và chuẩn đầu ra."),
            ("Lớp học", "Lớp Mầm/Chồi/Lá; thường học ban ngày, theo năm học.", "Lớp A1/A2/B1 hoặc giao tiếp; học theo ca và thời lượng khóa."),
            ("Giáo viên", "Chuyên môn giáo dục mầm non; theo dõi chăm sóc và phát triển trẻ.", "Chuyên môn ngoại ngữ; theo dõi kỹ năng nghe, nói, đọc, viết."),
            ("Tuyển sinh", "Ghi nhận độ tuổi, sức khỏe, người đón trẻ và ngày nhập học.", "Ghi nhận nhu cầu, trình độ đầu vào, lịch học phù hợp và mục tiêu."),
            ("Tài chính", "Học phí tháng, tiền ăn, bán trú và dịch vụ.", "Học phí khóa, giáo trình, lệ phí thi và học bổ trợ."),
            ("Kết quả", "Nhận xét theo 5 lĩnh vực phát triển.", "Điểm, xếp loại, kết quả thi và chứng chỉ."),
        ],
        [1.25, 2.65, 2.6],
    )
    base.add_callout(
        doc,
        "Nguyên tắc thực hiện:",
        "Tạo cùng một nhóm chức năng cho cả hai đơn vị, nhưng dùng bộ dữ liệu phù hợp chuyên ngành. "
        "Quy trình thao tác chỉ mô tả một lần; bảng trên nêu rõ dữ liệu cần thay đổi.",
        base.GREEN,
    )

    doc.add_heading("3. Hướng dẫn thao tác tuần tự", level=1)

    procedure(
        doc, 1, "Đăng nhập hệ thống",
        "Tất cả người dùng.",
        "Đã được cấp tên đăng nhập và mật khẩu.",
        "Xác thực người dùng và đưa người dùng đến đúng không gian làm việc.",
        [
            ("Mở địa chỉ http://localhost:5173.", "Hiển thị màn hình Đăng nhập."),
            ("Nhập tên đăng nhập.", "Tên đăng nhập được ghi nhận."),
            ("Nhập mật khẩu và chọn Đăng nhập.", "Hệ thống mở Portal đúng vai trò."),
        ],
        "00-dang-nhap.png",
        ["Tên người dùng và vai trò xuất hiện ở góc phải.", "Đơn vị đang làm việc hiển thị đúng.", "Menu bên trái phù hợp với quyền."],
        [
            ("TC-01-01", "Đăng nhập đúng.", "Mở đúng Portal."),
            ("TC-01-02", "Nhập sai mật khẩu.", "Báo lỗi và không tạo phiên."),
        ],
    )

    procedure(
        doc, 2, "Kiểm tra Portal quản lý đơn vị",
        "Quản lý đơn vị.",
        "Tài khoản demo_quanly_mn hoặc demo_quanly_nn.",
        "Rà soát số liệu tổng quan trước khi tạo dữ liệu để tránh thao tác nhầm đơn vị.",
        [
            ("Đăng nhập tài khoản quản lý.", "Mở Cổng quản lý đơn vị."),
            ("Kiểm tra tên đơn vị trên thanh trên cùng.", "Đúng trường/trung tâm cần triển khai."),
            ("Kiểm tra học viên, lớp, công nợ và việc chờ xử lý.", "Có số liệu nền hoặc bằng 0 nếu đơn vị mới."),
        ],
        "M01-01-portal-quan-ly.png",
        ["Đúng đơn vị.", "Không hiển thị dữ liệu của đơn vị khác.", "Các lối vào nhanh hoạt động."],
        [("TC-02-01", "Đăng nhập quản lý mầm non.", "Chỉ thấy dữ liệu Trường Mầm non Hoa Nắng.")],
    )

    procedure(
        doc, 3, "Tạo tài khoản nhân sự tại đơn vị",
        "Quản lý đơn vị hoặc quản trị có quyền quản lý người dùng.",
        "Đơn vị đã được tạo.",
        "Cấp tài khoản cho quản lý, tuyển sinh, tư vấn, học vụ và kế toán.",
        [
            ("Mở Hệ thống → Quản lý người dùng.", "Hiển thị biểu mẫu Tạo tài khoản mới."),
            ("Nhập tên đăng nhập, họ tên, email và số điện thoại.", "Thông tin hợp lệ."),
            ("Chọn vai trò ban đầu.", "Chỉ có Quản lý, Tuyển sinh, Tư vấn, Học vụ và Kế toán."),
            ("Chọn Tạo tài khoản.", "Hệ thống tạo mật khẩu tạm và thêm tài khoản vào danh sách."),
            ("Lặp lại bước mẫu cho các vai trò còn lại.", "Đủ nhân sự cần dùng trong đơn vị."),
        ],
        "M01-02-tao-tai-khoan-nhan-su.png",
        ["Không có lựa chọn Giáo viên hoặc Phụ huynh.", "Tài khoản gắn đúng đơn vị.", "Tên đăng nhập không trùng."],
        [
            ("TC-03-01", "Tạo tài khoản kế toán.", "Đăng nhập được Cổng kế toán."),
            ("TC-03-02", "Tạo tên đăng nhập đã tồn tại.", "Bị chặn."),
        ],
        "Giáo viên phải tạo từ hồ sơ Giáo viên; phụ huynh phải tạo từ liên kết phụ huynh - học sinh.",
    )

    procedure(
        doc, 4, "Tạo hồ sơ giáo viên",
        "Học vụ hoặc Quản lý đơn vị.",
        "Đơn vị đã có chương trình nhân sự.",
        "Tạo hồ sơ chuyên môn trước khi cấp tài khoản giáo viên.",
        [
            ("Mở Đào tạo → Giáo viên.", "Hiển thị danh sách giáo viên."),
            ("Chọn Thêm giáo viên.", "Mở biểu mẫu."),
            ("Nhập họ tên, điện thoại, email, chuyên môn và trình độ.", "Dữ liệu hợp lệ."),
            ("Chọn Tạo hồ sơ.", "Sinh mã giáo viên và thêm vào danh sách."),
        ],
        "M02-01-tao-ho-so-giao-vien.png",
        ["Mã giáo viên tự sinh.", "Hồ sơ thuộc đúng đơn vị.", "Trạng thái Hoạt động."],
        [("TC-04-01", "Tạo giáo viên đủ họ tên.", "Hồ sơ được tạo và có thể phân công lớp.")],
    )

    procedure(
        doc, 5, "Liên kết và cấp tài khoản giáo viên",
        "Học vụ hoặc Quản lý đơn vị.",
        "Hồ sơ giáo viên chưa có tài khoản.",
        "Liên kết tài khoản đăng nhập với đúng hồ sơ chuyên môn.",
        [
            ("Mở chi tiết giáo viên.", "Hiển thị khối Tài khoản đăng nhập."),
            ("Kiểm tra họ tên giáo viên đã đúng chính tả.", "Họ và tên là nguồn sinh tên đăng nhập."),
            ("Chọn Tạo tài khoản đăng nhập.", "Hệ thống sinh tên theo dạng gv_ten.ho và mật khẩu tạm."),
            ("Nếu tên đăng nhập đã tồn tại, kiểm tra số thứ tự.", "Hệ thống dùng gv_ten1.ho, gv_ten2.ho...; không dùng số điện thoại."),
            ("Bàn giao thông tin đăng nhập cho giáo viên.", "Giáo viên có thể đăng nhập."),
            ("Đăng nhập thử.", "Portal chỉ hiển thị lớp được phân công."),
        ],
        "NEW-01-quy-tac-tai-khoan-giao-vien.png",
        ["Tên đăng nhập không chứa số điện thoại.", "Không thể tạo tài khoản lần hai.", "Giáo viên không thấy lớp ngoài phân công."],
        [
            ("TC-05-01", "Tạo tài khoản cho Nguyễn Văn Phúc.", "Sinh gv_phuc.nguyen và liên kết đúng hồ sơ."),
            ("TC-05-02", "Tạo tài khoản cho giáo viên trùng tên.", "Tự thêm số sau phần tên; không báo trùng."),
        ],
        "Ví dụ: Nguyễn Văn Phúc → gv_phuc.nguyen. Chỉ bàn giao mật khẩu tạm qua kênh nội bộ an toàn.",
    )
    base.add_figure(
        doc,
        "NEW-02-ket-qua-tao-tai-khoan-giao-vien.png",
        "Hình 5.2. Kết quả sinh tài khoản từ họ tên giáo viên",
    )

    procedure(
        doc, 6, "Tiếp nhận và chăm sóc tuyển sinh",
        "Nhân viên Tuyển sinh hoặc Tư vấn.",
        "Đã đăng nhập tại đúng đơn vị.",
        "Quản lý người quan tâm từ lúc tiếp nhận đến trước khi xác nhận đăng ký.",
        [
            ("Mở Tuyển sinh → Tạo khách hàng tiềm năng.", "Mở biểu mẫu lead."),
            ("Nhập người liên hệ, điện thoại, nguồn, độ tuổi/trình độ và nhu cầu.", "Lead ở trạng thái Mới."),
            ("Mở chi tiết lead và ghi cuộc gọi/tin nhắn/lịch hẹn.", "Lịch sử chăm sóc được lưu."),
            ("Cập nhật Đang chăm sóc, Đã hẹn lịch hoặc Đã học thử.", "Trạng thái phản ánh đúng tiến độ."),
        ],
        "M03-01-tuyen-sinh.png",
        ["Lead có mã tự sinh.", "Có lịch sử liên hệ.", "Dữ liệu không lọt sang đơn vị khác."],
        [
            ("TC-06-01", "Tạo lead mới.", "Trạng thái Mới."),
            ("TC-06-02", "Đánh dấu Không tiếp tục thiếu lý do.", "Bị chặn."),
        ],
    )

    procedure(
        doc, 7, "Xác nhận đăng ký và hoàn thiện hồ sơ học sinh",
        "Tuyển sinh xác nhận; Học vụ hoàn thiện.",
        "Lead đủ điều kiện đăng ký.",
        "Chuyển người quan tâm thành học sinh/học viên và liên kết phụ huynh.",
        [
            ("Từ chi tiết lead chọn Xác nhận đăng ký.", "Mở biểu mẫu học sinh."),
            ("Nhập họ tên, ngày sinh, giới tính, địa chỉ và ngày nhập học.", "Tạo học sinh trạng thái Tiếp nhận."),
            ("Chọn quan hệ phụ huynh.", "Tạo hoặc tái sử dụng phụ huynh theo số điện thoại."),
            ("Kiểm tra form chi tiết nhập học.", "Rà soát học sinh, ngày nhập học, phụ huynh và người lập phiếu."),
            ("Chọn In phiếu.", "Mở bản in Phiếu xác nhận nhập học để ký và lưu hồ sơ."),
            ("Mở Học sinh · Học viên và chọn hồ sơ vừa tạo.", "Hiển thị hồ sơ chi tiết."),
            ("Bổ sung y tế, liên hệ khẩn cấp và thông tin chuyên ngành.", "Hồ sơ đầy đủ."),
        ],
        "NEW-06-phieu-xac-nhan-nhap-hoc.png",
        ["Học sinh chưa tự động thành Đang học.", "Phiếu ghi đúng đơn vị và ngày nhập học.", "Không tạo trùng khi xác nhận lại lead."],
        [
            ("TC-07-01", "Xác nhận lead.", "Tạo học sinh Tiếp nhận và phụ huynh chính."),
            ("TC-07-02", "Mở lại phiếu nhập học đã lập.", "Dữ liệu chứng từ không thay đổi và có thể in lại."),
        ],
    )

    procedure(
        doc, 8, "Tạo chương trình và lớp học",
        "Học vụ.",
        "Đã có giáo viên và kế hoạch đào tạo.",
        "Tạo khung chương trình và lớp thực tế để tiếp nhận học sinh.",
        [
            ("Mở Chương trình đào tạo → Tạo mới.", "Nhập tên, cấp độ, số buổi/giờ và mô tả."),
            ("Mở Lớp học → Tạo lớp.", "Chọn chương trình."),
            ("Nhập tên lớp, ngày bắt đầu/kết thúc, sĩ số và phòng.", "Lớp ở trạng thái Chuẩn bị."),
            ("Mở chi tiết lớp và chuyển Đang học khi sẵn sàng.", "Lớp có thể nhận học sinh."),
        ],
        "M05-01-lop-hoc.png",
        ["Lớp gắn đúng chương trình.", "Ngày và sĩ số hợp lệ.", "Không trộn lớp giữa các đơn vị."],
        [("TC-08-01", "Tạo lớp thiếu chương trình.", "Bị chặn.")],
        "Mầm non tổ chức theo độ tuổi và thường học cả ngày; ngoại ngữ tổ chức theo cấp độ/khóa và ca học.",
    )

    procedure(
        doc, 9, "Phân công giáo viên và xếp lớp",
        "Học vụ; Quản lý xử lý vượt sĩ số.",
        "Lớp Đang học, giáo viên Hoạt động, học sinh Tiếp nhận.",
        "Tạo quan hệ chính thức giữa lớp, giáo viên và học sinh.",
        [
            ("Mở chi tiết lớp → Phân công giáo viên.", "Chọn giáo viên, vai trò và ngày hiệu lực."),
            ("Chọn Xếp học sinh.", "Danh sách học sinh đủ điều kiện xuất hiện."),
            ("Chọn học sinh và ngày vào lớp.", "Tạo lượt xếp lớp Đang học."),
            ("Chọn Kèm phiếu xác nhận nhập học khi cần.", "Hệ thống lập đồng thời phiếu nhập học."),
            ("Mở In phiếu xếp lớp/In phiếu nhập học.", "Kiểm tra bản in trước khi ký."),
            ("Kiểm tra hồ sơ học sinh.", "Trạng thái tổng chuyển Đang học."),
        ],
        "M06-01-xep-lop-phan-cong.png",
        ["Giáo viên thấy lớp trên Portal.", "Sĩ số tăng đúng.", "Học sinh có một lượt active hợp lệ."],
        [("TC-09-01", "Xếp lớp đủ sĩ số.", "Người thường bị chặn; quản lý có quyền xử lý.")],
    )

    procedure(
        doc, 10, "Lập thời khóa biểu và sinh buổi học",
        "Học vụ.",
        "Lớp Đang học và đã phân công giáo viên.",
        "Sinh lịch học cụ thể để giáo viên có thể điểm danh và báo giảng.",
        [
            ("Mở Lịch học → Tạo quy tắc.", "Chọn lớp, thứ, giờ, phòng và giáo viên."),
            ("Nhập ngày áp dụng từ ngày hiện tại.", "Quy tắc có hiệu lực."),
            ("Chọn Sinh buổi học đến ngày.", "Các buổi học xuất hiện trong lịch."),
            ("Lọc theo khoảng ngày và giáo viên.", "Kiểm tra lịch không trùng."),
        ],
        "M07-01-thoi-khoa-bieu.png",
        ["Buổi học đúng thứ và khung giờ.", "Chạy sinh lại không tạo trùng.", "Không vượt ngày kết thúc lớp."],
        [("TC-10-01", "Sinh lịch 28 ngày hai lần.", "Số buổi không tăng ở lần hai.")],
    )

    procedure(
        doc, 11, "Bắt đầu buổi học và điểm danh",
        "Giáo viên hoặc Học vụ.",
        "Có buổi học hôm nay và roster học sinh.",
        "Ghi nhận tình trạng tham gia của từng học sinh.",
        [
            ("Mở Điểm danh và chọn buổi hôm nay.", "Hiển thị danh sách lớp."),
            ("Chọn Bắt đầu buổi học.", "Buổi chuyển Đang học."),
            ("Chọn Có mặt/Vắng có phép/Vắng không phép/Đi trễ/Về sớm.", "Trạng thái từng học sinh được chọn."),
            ("Nhập ghi chú/nhận xét và Lưu.", "Điểm danh được ghi nhận."),
            ("Ghi báo giảng và kết thúc buổi học.", "Portal giáo viên cập nhật."),
        ],
        "M07-02-diem-danh.png",
        ["Chỉ học sinh thuộc roster được lưu.", "Không điểm danh khi buổi chưa bắt đầu.", "Dữ liệu hiển thị cho phụ huynh phù hợp."],
        [("TC-11-01", "Điểm danh trước Bắt đầu.", "Bị chặn.")],
    )

    procedure(
        doc, 12, "Tiếp nhận và xử lý đơn xin phép",
        "Phụ huynh gửi; Học vụ/Giáo viên xử lý.",
        "Học sinh có lớp active.",
        "Quản lý xin nghỉ bằng trạng thái rõ ràng và có phản hồi.",
        [
            ("Phụ huynh mở tab Xin phép → Gửi đơn.", "Nhập lớp, thời gian và lý do."),
            ("Học vụ mở Đơn xin phép.", "Đơn ở trạng thái Chờ duyệt."),
            ("Chọn Duyệt hoặc Từ chối và nhập ghi chú.", "Trạng thái cập nhật."),
            ("Phụ huynh tải lại Portal.", "Thấy kết quả xử lý."),
        ],
        "M08-01-don-xin-phep.png",
        ["Không chọn được lớp đã kết thúc.", "Ghi chú xử lý hiển thị cho phụ huynh."],
        [("TC-12-01", "Từ chối đơn.", "Portal hiển thị Từ chối và ghi chú.")],
    )

    procedure(
        doc, 13, "Ghi nhận trao đổi với phụ huynh",
        "Giáo viên hoặc Học vụ.",
        "Có học sinh và phụ huynh liên kết.",
        "Lưu lịch sử liên lạc giữa nhà trường và gia đình.",
        [
            ("Mở Trao đổi phụ huynh → Ghi nhận mới.", "Mở biểu mẫu."),
            ("Chọn lớp/học sinh và kênh liên lạc.", "Đúng đối tượng."),
            ("Nhập nội dung, kết quả và thời gian.", "Lưu trao đổi."),
            ("Kiểm tra Portal phụ huynh.", "Trao đổi hiển thị đúng người ghi nhận."),
        ],
        "M08-02-trao-doi-phu-huynh.png",
        ["Không lộ trao đổi của học sinh khác.", "Vai trò hiển thị bằng tiếng Việt."],
        [("TC-13-01", "Giáo viên ghi trao đổi.", "Phụ huynh thấy nhãn Giáo viên.")],
    )

    procedure(
        doc, 14, "Tạo và gửi thông báo",
        "Quản lý hoặc Học vụ.",
        "Có đơn vị; phạm vi lớp/cá nhân cần lớp hoặc học sinh.",
        "Gửi thông tin đúng phạm vi Toàn trường, Theo lớp hoặc Cá nhân.",
        [
            ("Mở Thông báo nội bộ → Tạo mới.", "Nhập tiêu đề và nội dung."),
            ("Chọn phạm vi.", "Hệ thống yêu cầu lớp/học sinh nếu cần."),
            ("Chọn Lưu.", "Thông báo có mã và đơn vị gửi."),
            ("Phụ huynh mở menu Thông báo.", "Nhận đúng thông báo."),
        ],
        "M09-01-thong-bao.png",
        ["Phụ huynh nhiều đơn vị nhận đủ nguồn.", "Theo lớp không gửi nhầm lớp."],
        [("TC-14-01", "Hai đơn vị gửi Toàn trường.", "Phụ huynh hai con nhận đủ hai thông báo.")],
    )

    procedure(
        doc, 15, "Tạo kỳ thu, công nợ và thu tiền",
        "Kế toán.",
        "Học sinh đã xếp lớp.",
        "Thiết lập khoản thu và ghi nhận thanh toán theo đúng chuỗi chứng từ.",
        [
            ("Tạo Danh mục khoản thu.", "Khai báo loại và số tiền mặc định."),
            ("Tạo Kỳ thu và áp dụng các khoản thu.", "Kỳ ở trạng thái Nháp."),
            ("Mở kỳ và sinh khoản phải thu theo lớp.", "Mỗi học sinh active có công nợ."),
            ("Tại khoản phải thu chọn Thu tiền.", "Mở form chi tiết thu tiền."),
            ("Kiểm tra học sinh, lớp, số còn phải thu; nhập số tiền, ngày, phương thức và nội dung.", "Thông tin thu tiền đầy đủ."),
            ("Chọn Xác nhận thu tiền.", "Tạo phiếu thu và cập nhật công nợ."),
            ("Chọn In phiếu thu hoặc mở Lịch sử thu để in lại.", "Phiếu có số chứng từ và dữ liệu bất biến."),
            ("Đối chiếu Chưa thu/Thu một phần/Đã thu đủ.", "Trạng thái đúng số tiền."),
        ],
        "NEW-05-form-thu-tien.png",
        ["Không sinh trùng công nợ.", "Số đã thu và còn lại chính xác.", "Phiếu thu có thể mở/in lại từ lịch sử."],
        [
            ("TC-15-01", "Thu một phần.", "Trạng thái Thu một phần và tạo số phiếu."),
            ("TC-15-02", "Thu lớn hơn số còn phải thu.", "Hệ thống chặn và không tạo phiếu."),
            ("TC-15-03", "Mở lại phiếu thu từ Lịch sử thu.", "Thông tin khớp giao dịch gốc."),
        ],
    )
    base.add_figure(
        doc,
        "NEW-04-phieu-thu-chi-tiet.png",
        "Hình 15.2. Phiếu thu học phí có thể in và lưu hồ sơ",
    )

    procedure(
        doc, 16, "Xử lý điều chỉnh và chi phí",
        "Kế toán đề xuất; Quản lý đơn vị duyệt.",
        "Có công nợ hoặc nhu cầu chi phí.",
        "Kiểm soát hoàn phí, chuyển phí, bảo lưu và chi phí bằng quy trình phê duyệt.",
        [
            ("Mở Yêu cầu điều chỉnh → Tạo mới.", "Chọn hoàn/chuyển/bảo lưu và nhập lý do."),
            ("Quản lý mở yêu cầu và phê duyệt/từ chối.", "Trạng thái cập nhật."),
            ("Mở Chi phí → Đề xuất chi.", "Nhập danh mục, số tiền và lý do."),
            ("Quản lý phê duyệt.", "Chi phí được ghi nhận vào báo cáo."),
        ],
        "M10-03-chi-phi.png",
        ["Yêu cầu chưa duyệt chưa tác động báo cáo.", "Có lịch sử người đề xuất và người duyệt."],
        [("TC-16-01", "Chi phí Chờ duyệt.", "Không cộng vào tổng chi đã ghi nhận.")],
    )

    procedure(
        doc, 17, "Cập nhật kết quả học tập, thi và chứng chỉ",
        "Giáo viên hoặc Học vụ.",
        "Học sinh có lượt xếp lớp.",
        "Ghi nhận kết quả theo chuyên ngành và công bố cho phụ huynh.",
        [
            ("Mở hồ sơ học sinh → Lịch sử học tập/Kết quả.", "Chọn lượt xếp lớp."),
            ("Nhập loại đánh giá, ngày, điểm/xếp loại và nhận xét.", "Kết quả được lưu."),
            ("Mầm non chọn lĩnh vực phát triển.", "Hiển thị đúng 5 lĩnh vực."),
            ("Ngoại ngữ nhập điểm thi và nhận xét kỹ năng.", "Điểm gắn đúng lớp."),
            ("Mở Thành tích → Thêm chứng chỉ.", "Nhập tên, ngày, nơi cấp và minh chứng."),
        ],
        "M11-01-ket-qua-hoc-tap.png",
        ["Portal hiển thị nhãn tiếng Việt.", "Kết quả gắn đúng enrollment.", "Chứng chỉ hiển thị đúng học sinh."],
        [("TC-17-01", "Đánh giá mầm non Theo tháng · Ngôn ngữ.", "Portal hiển thị đúng nhãn.")],
    )

    procedure(
        doc, 18, "Kiểm tra Portal từng vai trò",
        "Quản lý, Tuyển sinh, Học vụ, Giáo viên, Kế toán và Phụ huynh.",
        "Đã hoàn thành dữ liệu từ Bước 1–17.",
        "Xác nhận mỗi vai trò chỉ nhìn thấy và thao tác đúng phạm vi.",
        [
            ("Đăng nhập lần lượt từng tài khoản demo.", "Về đúng Portal."),
            ("Đối chiếu menu và số liệu tổng quan.", "Khớp dữ liệu đã tạo."),
            ("Thử mở URL ngoài quyền.", "Bị chặn/điều hướng."),
            ("Đổi đơn vị nếu tài khoản có nhiều phạm vi.", "Dữ liệu tải lại theo đơn vị."),
        ],
        "05-portal-giao-vien-mam-non.png",
        ["Giáo viên chỉ thấy lớp phân công.", "Kế toán không sửa lớp.", "Phụ huynh chỉ thấy đúng con."],
        [("TC-18-01", "Giáo viên mở /finance.", "Bị chặn.")],
    )

    procedure(
        doc, 19, "Kiểm tra Portal phụ huynh đa đơn vị",
        "Phụ huynh.",
        "Một tài khoản liên kết hai con tại hai đơn vị.",
        "Xác nhận tổng hợp con, lịch, học phí, kết quả và thông báo theo từng đơn vị.",
        [
            ("Đăng nhập tài khoản phụ huynh.", "Tổng quan hiển thị số con/lớp."),
            ("Mở Thông tin các con.", "Dữ liệu được nhóm theo đơn vị."),
            ("Mở các tab Học tập, Học phí, Xin phép và Trao đổi.", "Đúng dữ liệu từng con."),
            ("Mở Thông báo.", "Hiển thị đủ thông báo từ các đơn vị."),
        ],
        "01-portal-phu-huynh-tong-quan.png",
        ["Hai con thuộc hai nhóm đơn vị.", "Không trộn lịch, kết quả hoặc công nợ.", "Thông báo ghi rõ đơn vị gửi."],
        [("TC-19-01", "Phụ huynh hai đơn vị đăng nhập.", "Nhìn thấy đủ hai con và hai nguồn thông báo.")],
    )

    procedure(
        doc, 20, "Đối chiếu báo cáo và bàn giao dữ liệu demo",
        "Quản lý đơn vị và Kế toán.",
        "Đã có kỳ thu, phiếu thu và chi phí.",
        "Đối chiếu số liệu tổng hợp trước khi bàn giao đào tạo hoặc demo.",
        [
            ("Mở Báo cáo tài chính.", "Chọn khoảng thời gian."),
            ("Đối chiếu tổng phải thu, đã thu, hoàn phí, chi phí và thu ròng.", "Khớp dữ liệu chi tiết."),
            ("Kiểm tra báo cáo theo đúng đơn vị.", "Không trộn số liệu."),
            ("Lưu biên bản test case.", "Đủ bằng chứng bàn giao."),
        ],
        "M12-01-bao-cao-tai-chinh.png",
        ["Số liệu báo cáo khớp công nợ và phiếu thu.", "Chi phí chỉ tính sau duyệt.", "Không truy cập đơn vị ngoài quyền."],
        [("TC-20-01", "Đối chiếu báo cáo mầm non.", "Khớp dữ liệu kỳ thu demo.")],
    )

    doc.add_heading("4. Hướng dẫn chi tiết theo nhóm chức năng", level=1)
    doc.add_paragraph(
        "Phần này dùng để đào tạo theo vai trò hoặc tra cứu một chức năng cụ thể. "
        "Mỗi nhóm đều mô tả riêng dữ liệu và cách vận hành tại trường mầm non và "
        "trung tâm ngoại ngữ. Các bước dùng chung được lặp lại có chủ đích để người "
        "đọc có thể thực hiện độc lập mà không phải suy đoán."
    )

    module_guide(
        doc,
        1,
        "Hệ thống, đơn vị, người dùng và phân quyền",
        "Quản trị hệ thống quản lý toàn hệ thống; Quản lý đơn vị quản lý nhân sự trong đơn vị được giao.",
        [
            ("Danh mục đơn vị", "Tạo trường, khối/lớp và thông tin liên hệ trường.", "Tạo trung tâm, cơ sở/địa điểm đào tạo và thông tin liên hệ."),
            ("Người dùng", "Tạo quản lý, tuyển sinh, học vụ, kế toán; giáo viên tạo từ hồ sơ.", "Tạo cùng nhóm vai trò; giáo viên tạo từ hồ sơ trung tâm."),
            ("Gán vai trò", "Gắn vai trò tại đúng trường; không dùng tài khoản trường khác.", "Gắn vai trò tại đúng trung tâm/cơ sở."),
            ("Vai trò · quyền", "Kiểm tra quyền xem/quản lý học sinh, lớp, điểm danh và tài chính.", "Kiểm tra quyền tuyển sinh, lịch theo ca, kết quả thi và tài chính khóa."),
            ("Nhật ký", "Tra cứu người tạo/sửa/xóa hoặc duyệt dữ liệu.", "Tra cứu tương tự, lọc đúng trung tâm."),
            ("Tìm kiếm", "Tìm học sinh, giáo viên, lớp theo tên/mã.", "Tìm học viên, giáo viên, lớp/khóa theo tên/mã."),
        ],
        [
            ("Đăng nhập tài khoản quản trị và mở Danh mục đơn vị.", "Chọn đúng Trường Mầm non Hoa Nắng."),
            ("Mở Quản lý người dùng tại trường.", "Tạo lần lượt quản lý, tuyển sinh, học vụ và kế toán."),
            ("Mở chi tiết từng tài khoản.", "Kiểm tra trạng thái, đơn vị, vai trò và quyền hiệu lực."),
            ("Mở Vai trò · Phân quyền.", "Đối chiếu quyền theo ma trận đã phê duyệt."),
            ("Mở Nhật ký hệ thống.", "Lọc theo hành động và tài khoản vừa tạo."),
        ],
        [
            ("Đăng nhập quản trị và chọn Trung tâm Ngoại ngữ Quận 8.", "Không còn dữ liệu của trường mầm non."),
            ("Tạo cùng nhóm tài khoản vận hành cho trung tâm.", "Tài khoản thuộc đúng trung tâm."),
            ("Kiểm tra quyền học vụ và tuyển sinh.", "Có quyền quản lý học viên/lớp nhưng không có quyền duyệt tài chính nếu chưa được giao."),
            ("Thử mở URL ngoài quyền bằng tài khoản giáo viên/kế toán.", "Hệ thống chặn hoặc điều hướng an toàn."),
            ("Tra cứu Nhật ký hệ thống.", "Mỗi thao tác ghi đúng người, đơn vị và thời điểm."),
        ],
        [
            ("TC-4.1-01", "Tạo người dùng tại trường rồi mở tại trung tâm.", "Không nhìn thấy nếu không được gán phạm vi."),
            ("TC-4.1-02", "Giáo viên mở Quản lý người dùng.", "Bị chặn."),
            ("TC-4.1-03", "Tìm kiếm theo mã học sinh/lớp.", "Trả về đúng dữ liệu thuộc đơn vị hiện tại."),
        ],
    )

    module_guide(
        doc,
        2,
        "Tuyển sinh và chăm sóc khách hàng tiềm năng",
        "Tuyển sinh/Tư vấn tiếp nhận và chăm sóc; Học vụ phối hợp xác nhận dữ liệu đào tạo.",
        [
            ("Tiếp nhận lead", "Ghi tên phụ huynh, điện thoại, độ tuổi của trẻ, nhu cầu bán trú.", "Ghi người liên hệ/học viên, trình độ quan tâm, mục tiêu và ca học."),
            ("Nguồn tuyển sinh", "Website, giới thiệu phụ huynh, sự kiện tại trường.", "Website, quảng cáo, giới thiệu, hội thảo hoặc học thử."),
            ("Hoạt động chăm sóc", "Cuộc gọi, tin nhắn, tham quan trường, hẹn trao đổi.", "Cuộc gọi, tư vấn, hẹn test đầu vào, học thử."),
            ("Trạng thái", "Mới → Đang chăm sóc → Đã hẹn → Đã đăng ký/Không tiếp tục.", "Mới → Đang chăm sóc → Đã hẹn → Đã học thử → Đã đăng ký/Không tiếp tục."),
            ("Xác nhận đăng ký", "Tạo hồ sơ trẻ và liên kết phụ huynh.", "Tạo hồ sơ học viên; có thể bổ sung kết quả test đầu vào."),
            ("Lịch hẹn", "Theo dõi lịch tham quan/trao đổi còn chờ xử lý.", "Theo dõi lịch tư vấn/test/học thử, kể cả lịch đã quá giờ."),
        ],
        [
            ("Mở Tuyển sinh → Tiếp nhận khách hàng tiềm năng.", "Nhập người liên hệ là phụ huynh."),
            ("Nhập điện thoại, nguồn, độ tuổi 4–5 và nhu cầu học bán trú.", "Lead ở trạng thái Mới."),
            ("Mở chi tiết lead, ghi cuộc gọi và đặt lịch tham quan.", "Lịch sử chăm sóc có thời gian và kết quả."),
            ("Sau khi phụ huynh đồng ý, chọn Xác nhận đăng ký.", "Nhập thông tin trẻ, ngày nhập học và quan hệ phụ huynh."),
            ("Kiểm tra hồ sơ học sinh.", "Trẻ ở trạng thái Tiếp nhận, chưa tự động vào lớp."),
        ],
        [
            ("Tạo lead có nhu cầu tiếng Anh A2.", "Ghi rõ mục tiêu giao tiếp/chứng chỉ và ca tối."),
            ("Ghi hoạt động tư vấn và đặt lịch test đầu vào.", "Trạng thái chuyển Đã hẹn lịch."),
            ("Ghi kết quả học thử hoặc test.", "Lịch sử hoạt động đầy đủ."),
            ("Chọn Xác nhận đăng ký.", "Tạo học viên và liên kết người giám hộ nếu học viên chưa thành niên."),
            ("Chuyển hồ sơ cho Học vụ.", "Học viên xuất hiện trong danh sách chờ xếp lớp."),
        ],
        [
            ("TC-4.2-01", "Tạo lead thiếu số điện thoại bắt buộc.", "Hệ thống báo dữ liệu chưa hợp lệ."),
            ("TC-4.2-02", "Đánh dấu Không tiếp tục không có lý do.", "Hệ thống chặn."),
            ("TC-4.2-03", "Xác nhận cùng lead lần hai.", "Không tạo trùng học sinh/học viên."),
            ("TC-4.2-04", "Lọc lịch hẹn sắp tới/quá hạn.", "Hiển thị đúng lịch chưa xử lý."),
        ],
    )

    module_guide(
        doc,
        3,
        "Hồ sơ học sinh, học viên và phụ huynh",
        "Tuyển sinh tạo hồ sơ ban đầu; Học vụ cập nhật hồ sơ đào tạo; Quản lý kiểm soát dữ liệu.",
        [
            ("Thông tin cơ bản", "Họ tên, tên gọi, ngày sinh, giới tính, ngày nhập học.", "Họ tên, ngày sinh, giới tính, ngày đăng ký/nhập học."),
            ("Định danh", "Giấy khai sinh/định danh, nơi sinh, dân tộc, quốc tịch, diện ưu tiên.", "CCCD/định danh, nơi sinh, quốc tịch, trường học trước đó."),
            ("Sức khỏe", "Chiều cao, cân nặng, dị ứng, bệnh nền, liên hệ khẩn cấp.", "Thông tin sức khỏe và liên hệ khẩn cấp khi cần."),
            ("Phụ huynh", "Bắt buộc theo dõi quan hệ, người đón trẻ, liên hệ chính.", "Liên kết người giám hộ với học viên chưa thành niên; học viên người lớn có thể tự liên hệ."),
            ("Trạng thái", "Tiếp nhận, Đang học, Bảo lưu, Ngừng học, Hoàn thành.", "Tiếp nhận, Đang học, Bảo lưu, Ngừng học, Hoàn thành."),
            ("Học tập/thành tích", "Lịch sử lớp, nhận xét 5 lĩnh vực, thành tích.", "Lịch sử khóa/lớp, điểm, xếp loại, chứng chỉ."),
        ],
        [
            ("Mở Học sinh · Học viên → Thêm học sinh.", "Nhập hồ sơ mẫu của một trẻ."),
            ("Mở chi tiết → Thông tin.", "Bổ sung định danh, sức khỏe và liên hệ khẩn cấp."),
            ("Mở Phụ huynh → Thêm phụ huynh.", "Nhập số điện thoại trước để tái sử dụng hồ sơ có sẵn."),
            ("Đặt người liên hệ chính và quan hệ.", "Chỉ một liên kết chính theo quy tắc nghiệp vụ."),
            ("Kiểm tra Lịch sử học tập, Thành tích và Kết quả.", "Các mục tách đúng tab, đúng trẻ."),
        ],
        [
            ("Tạo hồ sơ học viên sau xác nhận tuyển sinh.", "Nhập thông tin cá nhân và trình độ ban đầu."),
            ("Liên kết phụ huynh/người giám hộ nếu cần.", "Không tạo hồ sơ phụ huynh trùng số điện thoại."),
            ("Ghi kết quả test đầu vào khi xếp lớp.", "Kết quả gắn với lượt xếp lớp."),
            ("Sau khóa học, thêm chứng chỉ/thành tích.", "Có tên, kết quả, ngày đạt, nơi cấp và minh chứng."),
            ("Chuyển trạng thái Hoàn thành.", "Lịch sử trạng thái giữ nguyên các mốc trước."),
        ],
        [
            ("TC-4.3-01", "Dùng cùng số điện thoại phụ huynh cho hai con.", "Tái sử dụng một tài khoản phụ huynh."),
            ("TC-4.3-02", "Phụ huynh đã có ở đơn vị khác.", "Hiện cảnh báo và liên kết có kiểm soát."),
            ("TC-4.3-03", "Gỡ liên kết phụ huynh chính.", "Yêu cầu xác nhận và không làm mất hồ sơ học sinh."),
            ("TC-4.3-04", "Đổi trạng thái có ngày hiệu lực.", "Lịch sử ghi đúng trạng thái, lý do và ngày."),
        ],
    )

    module_guide(
        doc,
        4,
        "Giáo viên và tài khoản giáo viên",
        "Học vụ/Quản lý tạo hồ sơ, cấp tài khoản và phân công; giáo viên cập nhật hồ sơ cá nhân.",
        [
            ("Hồ sơ chuyên môn", "Chuyên môn chăm sóc–giáo dục trẻ, trình độ sư phạm mầm non.", "Chuyên môn ngoại ngữ, kỹ năng/chuẩn chứng chỉ, trình độ sư phạm."),
            ("Tài khoản", "Tạo từ hồ sơ; tên dạng gv_ten.ho, trùng thì thêm số.", "Áp dụng cùng quy tắc; không dùng số điện thoại."),
            ("Trạng thái", "Hoạt động/Ngừng hoạt động; chỉ giáo viên hoạt động được phân công.", "Áp dụng tương tự."),
            ("Phân công", "Giáo viên chính/phụ, thường theo lớp trong năm học.", "Giáo viên chính/trợ giảng, có thể theo lớp/ca/khóa."),
            ("Portal", "Lớp, lịch, điểm danh, báo giảng, nhận xét phát triển.", "Lớp, lịch, điểm danh, báo giảng, điểm/xếp loại/chứng chỉ."),
        ],
        [
            ("Mở Giáo viên → Thêm giáo viên.", "Nhập họ tên, liên hệ, chuyên môn và trình độ."),
            ("Mở chi tiết hồ sơ.", "Kiểm tra mã giáo viên tự sinh và trạng thái Hoạt động."),
            ("Chọn Tạo tài khoản đăng nhập.", "Nhận tên gv_ten.ho và mật khẩu tạm."),
            ("Phân công vào lớp Mầm/Chồi/Lá.", "Chọn vai trò và ngày hiệu lực."),
            ("Đăng nhập Portal giáo viên.", "Chỉ thấy lớp mầm non được phân công."),
        ],
        [
            ("Tạo giáo viên tiếng Anh.", "Ghi chuyên môn, cấp độ giảng dạy và trình độ."),
            ("Tạo tài khoản từ hồ sơ.", "Không nhập tên đăng nhập thủ công từ số điện thoại."),
            ("Phân công vào lớp A2 hoặc lớp giao tiếp.", "Chọn giáo viên chính/trợ giảng."),
            ("Kiểm tra xung đột lịch.", "Không phân công khung giờ trùng không hợp lệ."),
            ("Đăng nhập Portal.", "Chỉ thấy lớp và lịch thuộc trung tâm."),
        ],
        [
            ("TC-4.4-01", "Nguyễn Văn Phúc chưa có tài khoản.", "Sinh gv_phuc.nguyen."),
            ("TC-4.4-02", "Giáo viên trùng họ tên.", "Sinh gv_phuc1.nguyen hoặc số tiếp theo."),
            ("TC-4.4-03", "Tạo tài khoản lần hai.", "Bị chặn."),
            ("TC-4.4-04", "Ngừng hoạt động giáo viên.", "Không cho phân công mới; dữ liệu lịch sử còn nguyên."),
        ],
    )

    module_guide(
        doc,
        5,
        "Chương trình đào tạo, lớp học và xếp lớp",
        "Học vụ cấu hình; Quản lý phê duyệt các ngoại lệ về sĩ số hoặc tổ chức.",
        [
            ("Chương trình", "Tên chương trình theo độ tuổi; số buổi/giờ và kế hoạch giáo dục.", "Tên khóa/cấp độ; tổng buổi, tổng giờ và chuẩn đầu ra."),
            ("Trạng thái chương trình", "Hoạt động/Ngừng hoạt động.", "Hoạt động/Ngừng hoạt động."),
            ("Lớp", "Tên lớp Mầm/Chồi/Lá, phòng, năm học và sĩ số.", "Tên lớp theo cấp độ/ca, phòng, ngày khóa học và sĩ số."),
            ("Giáo viên", "Phân công giáo viên chính/phụ.", "Phân công giáo viên chính/trợ giảng."),
            ("Xếp lớp", "Theo độ tuổi, khả năng chăm sóc và sĩ số.", "Theo nguyện vọng, kết quả test, cấp độ và ca học."),
            ("Điều chuyển", "Chuyển lớp, ngừng học, hoàn thành; lưu ngày/lý do.", "Chuyển cấp độ/ca/lớp; bảo toàn lịch sử khóa."),
            ("Phiếu", "Phiếu xếp lớp và phiếu xác nhận nhập học.", "Áp dụng cùng mẫu, mang thông tin trung tâm/khóa."),
        ],
        [
            ("Tạo chương trình cho nhóm tuổi 5–6.", "Nhập số buổi/giờ và mô tả."),
            ("Tạo lớp Lá, chọn chương trình, ngày học, sĩ số và phòng.", "Lớp ở trạng thái Chuẩn bị."),
            ("Phân công giáo viên.", "Chọn vai trò và ngày hiệu lực."),
            ("Chọn trẻ chờ xếp lớp, nhập ngày vào lớp.", "Có thể chọn kèm phiếu nhập học."),
            ("Chuyển lớp sang Đang học.", "Sĩ số và hồ sơ trẻ cập nhật."),
        ],
        [
            ("Tạo chương trình A2 hoặc giao tiếp.", "Nhập tổng buổi, tổng giờ và mô tả chuẩn đầu ra."),
            ("Tạo lớp theo ca tối/cuối tuần.", "Ngày bắt đầu–kết thúc phù hợp khóa."),
            ("Xếp học viên theo kết quả test đầu vào.", "Lưu kết quả test và ghi chú."),
            ("Phân công giáo viên đúng chuyên môn.", "Lớp xuất hiện trên Portal giáo viên."),
            ("Kết thúc khóa.", "Chuyển lượt xếp lớp Hoàn thành và giữ lịch sử."),
        ],
        [
            ("TC-4.5-01", "Xếp vượt sĩ số.", "Người dùng thường bị chặn; quản lý xử lý theo quyền."),
            ("TC-4.5-02", "Một học sinh có hai lượt active xung đột.", "Hệ thống chặn hoặc yêu cầu chuyển lớp."),
            ("TC-4.5-03", "Ngừng chương trình đang có lớp.", "Lớp cũ vẫn xem được; không dùng cho lớp mới."),
            ("TC-4.5-04", "In lại phiếu xếp lớp.", "Dữ liệu chứng từ giữ nguyên."),
        ],
    )

    module_guide(
        doc,
        6,
        "Lịch học, buổi học, điểm danh và báo giảng",
        "Học vụ lập lịch; Giáo viên vận hành buổi học; Quản lý theo dõi ngoại lệ.",
        [
            ("Khung giờ", "Thường theo lịch ngày trong tuần, phòng/lớp cố định.", "Theo ca tối/cuối tuần, có thể nhiều khung trong một lớp."),
            ("Sinh buổi học", "Sinh đến ngày trong năm học, không trùng.", "Sinh theo thời lượng khóa, không vượt ngày kết thúc."),
            ("Điều chỉnh", "Nghỉ/hủy, xếp bù, đổi phòng/giáo viên.", "Nghỉ/hủy, học bù, đổi ca/phòng/giáo viên."),
            ("Điểm danh", "Có mặt, vắng phép, vắng không phép, đi trễ, về sớm.", "Áp dụng cùng trạng thái."),
            ("Báo giảng", "Nội dung hoạt động, bài tập và ghi chú chăm sóc.", "Nội dung bài học, bài tập và ghi chú chuyên môn."),
            ("Bộ lọc", "Theo ngày, lớp và giáo viên.", "Theo khoảng ngày, ca/lớp và giáo viên."),
        ],
        [
            ("Trong chi tiết lớp chọn Thêm khung giờ.", "Chọn thứ, giờ, phòng và giáo viên."),
            ("Chọn Sinh buổi học đến ngày.", "Kiểm tra các buổi xuất hiện trên Lịch học."),
            ("Mở buổi hôm nay và chọn Bắt đầu.", "Buổi chuyển Đang học."),
            ("Điểm danh toàn bộ roster.", "Ghi chú trẻ đi trễ/về sớm khi cần."),
            ("Nhập báo giảng và Kết thúc buổi.", "Buổi hoàn thành, phụ huynh xem được dữ liệu phù hợp."),
        ],
        [
            ("Tạo lịch thứ 2–4–6, ca 18:00–19:30.", "Gắn lớp A2 và giáo viên."),
            ("Sinh buổi đến ngày kết thúc khóa.", "Không tạo sau ngày kết thúc."),
            ("Lọc lịch theo giáo viên.", "Phát hiện ca trùng."),
            ("Bắt đầu, điểm danh và ghi bài học/bài tập.", "Dữ liệu gắn đúng buổi."),
            ("Hủy một buổi và tạo lịch bù.", "Buổi gốc giữ trạng thái; buổi bù có thời gian mới."),
        ],
        [
            ("TC-4.6-01", "Sinh cùng khoảng ngày hai lần.", "Không tạo buổi trùng."),
            ("TC-4.6-02", "Điểm danh trước khi bắt đầu.", "Bị chặn."),
            ("TC-4.6-03", "Kết thúc khi chưa lưu roster.", "Cảnh báo hoặc yêu cầu hoàn tất."),
            ("TC-4.6-04", "Lịch vượt ngày kết thúc lớp.", "Không được sinh."),
        ],
    )

    module_guide(
        doc,
        7,
        "Portal giáo viên: công việc hằng ngày",
        "Giáo viên thực hiện trong phạm vi lớp được phân công; Học vụ hỗ trợ điều phối.",
        [
            ("Tổng quan", "Số lớp, buổi sắp tới, việc cần điểm danh.", "Số lớp/ca, buổi sắp tới và việc cần xử lý."),
            ("Thông báo", "Đọc thông báo vận hành, lịch sự kiện, nhắc việc.", "Đọc thông báo đổi lịch, thi, học bù, vận hành."),
            ("Điểm danh", "Thực hiện theo buổi, ghi trạng thái từng trẻ.", "Thực hiện theo ca, ghi trạng thái từng học viên."),
            ("Kết quả", "Nhận xét theo tháng/quý/năm và 5 lĩnh vực phát triển.", "Nhập điểm, xếp loại, nhận xét kỹ năng, kết quả thi."),
            ("Xin phép", "Theo dõi đơn phụ huynh gửi; xử lý nếu được phân quyền.", "Theo dõi đơn nghỉ/học bù; xử lý nếu được phân quyền."),
            ("Trao đổi", "Ghi nhận gọi điện/tin nhắn/gặp trực tiếp với phụ huynh.", "Ghi trao đổi kết quả học, chuyên cần và kế hoạch học."),
            ("Báo giảng", "Hoạt động học/chăm sóc, ghi chú và bài tập.", "Nội dung bài học, bài tập và ghi chú."),
            ("Hồ sơ cá nhân", "Cập nhật ảnh, họ tên, email, điện thoại; đổi mật khẩu khi được yêu cầu.", "Áp dụng tương tự."),
        ],
        [
            ("Đăng nhập Portal giáo viên mầm non.", "Kiểm tra tên trường và danh sách lớp."),
            ("Mở Thông báo.", "Đọc thông báo mới; kiểm tra đơn vị gửi và đính kèm."),
            ("Mở Lịch học/Điểm danh.", "Chọn đúng buổi, bắt đầu và ghi roster."),
            ("Mở lớp → Ghi kết quả học tập.", "Chọn Theo tháng và một trong 5 lĩnh vực; nhập nhận xét."),
            ("Mở Đơn xin phép.", "Xem lý do, thời gian; duyệt/từ chối nếu có quyền."),
            ("Mở Trao đổi phụ huynh.", "Chọn trẻ, kênh, nội dung và hướng xử lý."),
            ("Hoàn tất Báo giảng.", "Nhập hoạt động, bài tập và ghi chú; kết thúc buổi."),
        ],
        [
            ("Đăng nhập Portal giáo viên ngoại ngữ.", "Kiểm tra lớp/ca đúng phân công."),
            ("Đọc thông báo đổi lịch, kỳ thi hoặc học bù.", "Thông báo được đánh dấu đã đọc theo giao diện."),
            ("Điểm danh buổi tối.", "Ghi vắng phép/không phép và đi trễ."),
            ("Ghi kết quả giữa kỳ/cuối kỳ.", "Nhập điểm, xếp loại, nhận xét nghe–nói–đọc–viết."),
            ("Theo dõi đơn xin nghỉ.", "Phối hợp học vụ sắp lịch bù nếu cần."),
            ("Ghi trao đổi với phụ huynh/học viên.", "Nội dung gắn đúng học viên/lớp."),
            ("Nhập báo giảng và bài tập.", "Portal phụ huynh/học viên nhận đúng thông tin được công bố."),
        ],
        [
            ("TC-4.7-01", "Giáo viên mở lớp không được phân công.", "Không nhìn thấy hoặc bị chặn."),
            ("TC-4.7-02", "Ghi đánh giá mầm non thiếu lĩnh vực.", "Yêu cầu bổ sung theo nghiệp vụ."),
            ("TC-4.7-03", "Ghi điểm ngoại ngữ ngoài lớp phụ trách.", "Bị chặn."),
            ("TC-4.7-04", "Mở /finance bằng giáo viên.", "Bị chặn."),
            ("TC-4.7-05", "Đơn xin phép đã xử lý.", "Không xử lý lần hai; giữ người và thời gian xử lý."),
        ],
    )
    base.add_callout(
        doc,
        "Phân định trách nhiệm về đơn xin phép:",
        "Theo chức năng hiện tại, phụ huynh là người gửi đơn xin nghỉ cho học sinh/học viên. "
        "Giáo viên theo dõi và xử lý khi được phân quyền, không tạo đơn thay phụ huynh. "
        "Nếu phụ huynh báo qua điện thoại, giáo viên ghi nhận tại Trao đổi phụ huynh và đề nghị "
        "Học vụ xử lý theo quy định của đơn vị.",
        base.AMBER,
    )

    module_guide(
        doc,
        8,
        "Thông báo, xin phép và trao đổi phụ huynh",
        "Quản lý/Học vụ tạo thông báo; Phụ huynh gửi đơn; Giáo viên/Học vụ xử lý và trao đổi.",
        [
            ("Thông báo toàn đơn vị", "Lịch nghỉ, sự kiện, thực đơn, họp phụ huynh.", "Lịch học, lịch thi, khai giảng, thay đổi ca."),
            ("Thông báo theo lớp", "Gửi riêng lớp Mầm/Chồi/Lá.", "Gửi riêng lớp/cấp độ/khóa."),
            ("Thông báo cá nhân", "Gửi cho một trẻ/phụ huynh.", "Gửi cho một học viên/người giám hộ."),
            ("Đính kèm", "Tên tệp và liên kết tài liệu/hình ảnh.", "Giáo trình, lịch thi, tài liệu học."),
            ("Đơn xin phép", "Phụ huynh chọn con/lớp, thời gian, lý do.", "Phụ huynh/học viên chọn lớp/ca, thời gian, lý do."),
            ("Trao đổi", "Gọi điện, tin nhắn, gặp trực tiếp; ghi kết quả.", "Trao đổi chuyên cần, kết quả, lịch học và lộ trình."),
        ],
        [
            ("Tạo thông báo Toàn trường.", "Nhập tiêu đề, nội dung và đính kèm nếu có."),
            ("Tạo thông báo Theo lớp.", "Chọn đúng lớp Lá."),
            ("Đăng nhập phụ huynh.", "Kiểm tra nhận đủ thông báo từ trường."),
            ("Gửi đơn xin phép cho một trẻ.", "Chọn thời gian và lý do."),
            ("Học vụ/Giáo viên xử lý.", "Phụ huynh thấy trạng thái và ghi chú."),
            ("Ghi trao đổi sau khi liên hệ.", "Lịch sử có người ghi, vai trò, kênh và kết quả."),
        ],
        [
            ("Tạo thông báo lịch thi toàn trung tâm.", "Gắn tên/liên kết tài liệu nếu cần."),
            ("Tạo thông báo đổi ca theo lớp A2.", "Chỉ học viên lớp A2 nhận."),
            ("Phụ huynh/học viên gửi đơn xin nghỉ.", "Đơn gắn đúng lớp và thời gian."),
            ("Học vụ xử lý và lên lịch bù.", "Kết quả xử lý rõ ràng."),
            ("Giáo viên ghi trao đổi về tiến độ học.", "Portal hiển thị đúng người ghi."),
        ],
        [
            ("TC-4.8-01", "Phụ huynh có con ở hai đơn vị.", "Nhận đủ hai thông báo toàn đơn vị, ghi rõ nguồn."),
            ("TC-4.8-02", "Thông báo Theo lớp thiếu lớp.", "Bị chặn."),
            ("TC-4.8-03", "Thông báo Cá nhân thiếu học sinh.", "Bị chặn."),
            ("TC-4.8-04", "Kế toán mở Trao đổi phụ huynh.", "Bị chặn nếu không có quyền."),
        ],
    )

    module_guide(
        doc,
        9,
        "Tài chính, công nợ, thu tiền, điều chỉnh và chi phí",
        "Kế toán lập và theo dõi; Quản lý duyệt điều chỉnh/chi phí theo cấu hình.",
        [
            ("Khoản thu", "Học phí tháng, tiền ăn, bán trú, dịch vụ.", "Học phí khóa, giáo trình, lệ phí thi, bổ trợ."),
            ("Kỳ thu", "Theo tháng/học kỳ/đợt.", "Theo khóa/học kỳ/đợt."),
            ("Công nợ", "Sinh theo lớp và học sinh đang học.", "Sinh theo lớp/khóa và học viên active."),
            ("Thu tiền", "Tiền mặt/chuyển khoản/thẻ/khác; thu một phần hoặc đủ.", "Áp dụng tương tự."),
            ("Miễn giảm", "Nhập số tiền giảm theo chính sách.", "Học bổng/ưu đãi/giảm phí theo chương trình."),
            ("Điều chỉnh", "Hoàn phí, chuyển phí, bảo lưu; có lý do và duyệt.", "Hoàn/chuyển sang khóa khác/bảo lưu."),
            ("Chi phí", "Thực phẩm, dịch vụ, mua sắm, vận hành.", "Giảng viên, phòng học, tài liệu, thi cử, vận hành."),
            ("Báo cáo", "Tổng thu, hoàn phí, công nợ, chi, lãi/lỗ.", "Áp dụng cùng chỉ số; lọc đúng trung tâm."),
        ],
        [
            ("Tạo danh mục Học phí, Tiền ăn và Dịch vụ.", "Nhập loại và số tiền mặc định."),
            ("Tạo kỳ thu tháng, thêm các khoản thu và mở kỳ.", "Kỳ không còn sửa thông tin nền sau khi mở."),
            ("Sinh công nợ theo lớp Lá.", "Mỗi trẻ active có một khoản phải thu."),
            ("Chọn Thu tiền, nhập số tiền/ngày/phương thức/nội dung.", "Tạo phiếu thu."),
            ("Nếu có miễn giảm/hoàn phí, tạo yêu cầu.", "Quản lý duyệt trước khi báo cáo thay đổi."),
            ("Tạo đề xuất chi và duyệt.", "Chi phí đã duyệt xuất hiện trong báo cáo."),
        ],
        [
            ("Tạo Học phí khóa A2 và Giáo trình.", "Nhập số tiền mặc định."),
            ("Tạo kỳ thu Theo khóa.", "Ngày và hạn thanh toán theo khóa."),
            ("Sinh công nợ lớp A2.", "Không sinh trùng khi chạy lại."),
            ("Thu một phần rồi thu phần còn lại.", "Hai phiếu, công nợ chuyển Thu một phần → Đã thu đủ."),
            ("Tạo yêu cầu chuyển phí sang khóa khác.", "Chọn khoản phải thu đích và lý do."),
            ("Đối chiếu báo cáo.", "Tổng thu/chi/công nợ khớp chứng từ."),
        ],
        [
            ("TC-4.9-01", "Thu lớn hơn số còn phải thu.", "Bị chặn."),
            ("TC-4.9-02", "Sinh công nợ cùng lớp hai lần.", "Bỏ qua khoản đã có."),
            ("TC-4.9-03", "Điều chỉnh chưa duyệt.", "Chưa tác động số liệu chính thức."),
            ("TC-4.9-04", "Chi phí Chờ duyệt.", "Không cộng vào tổng chi đã ghi nhận."),
            ("TC-4.9-05", "Xuất CSV báo cáo.", "Tệp có đúng kỳ thu và số liệu đang lọc."),
        ],
    )

    module_guide(
        doc,
        10,
        "Phiếu nhập học, phiếu xếp lớp và phiếu thu",
        "Học vụ lập phiếu đào tạo; Kế toán lập phiếu thu; Quản lý cấu hình mẫu in.",
        [
            ("Phiếu nhập học", "Tên trường, trẻ, ngày nhập học, phụ huynh/người lập.", "Tên trung tâm, học viên, ngày nhập học, người lập."),
            ("Phiếu xếp lớp", "Lớp theo độ tuổi, ngày vào lớp, người xếp.", "Lớp/cấp độ, kết quả test, ngày vào lớp."),
            ("Phiếu thu", "Trẻ, lớp, kỳ thu, số tiền, phương thức, bằng chữ.", "Học viên, lớp/khóa, kỳ thu, số tiền, phương thức, bằng chữ."),
            ("Mẫu in", "Logo, nhãn chữ ký người lập/người nộp/đại diện, footer.", "Cấu hình riêng trung tâm, không dùng nhầm mẫu trường."),
            ("In lại", "Mở chứng từ cũ từ hồ sơ/lịch sử thu.", "Áp dụng tương tự."),
        ],
        [
            ("Tại bước xếp lớp chọn Kèm phiếu xác nhận nhập học.", "Hệ thống sinh số phiếu."),
            ("Mở Phiếu nhập học và Phiếu xếp lớp.", "Kiểm tra tên trường, trẻ, lớp và ngày."),
            ("Tại kỳ thu, ghi nhận một giao dịch.", "Hệ thống sinh Phiếu thu."),
            ("Mở Lịch sử thu và chọn số phiếu.", "Mở đúng chứng từ cũ."),
            ("Chọn In.", "Hộp thoại in hiển thị mẫu không có menu thao tác."),
        ],
        [
            ("Lập phiếu nhập học cho học viên trung tâm.", "Kiểm tra tên trung tâm và học viên."),
            ("Xếp lớp theo kết quả test và mở phiếu xếp lớp.", "Phiếu ghi đúng lớp/cấp độ."),
            ("Ghi nhận học phí khóa và mở phiếu thu.", "Số tiền và bằng chữ khớp."),
            ("Mở Cài đặt → Thiết lập mẫu in.", "Cấu hình logo, nhãn ký và footer riêng."),
            ("In lại chứng từ.", "Dữ liệu không bị cập nhật theo hồ sơ thay đổi sau này."),
        ],
        [
            ("TC-4.10-01", "Mở phiếu không thuộc đơn vị.", "Bị chặn/không tìm thấy."),
            ("TC-4.10-02", "In lại phiếu thu.", "Số phiếu và dữ liệu khớp giao dịch gốc."),
            ("TC-4.10-03", "Đổi cấu hình mẫu trung tâm.", "Không ảnh hưởng mẫu của trường mầm non."),
        ],
    )

    module_guide(
        doc,
        11,
        "Portal phụ huynh và dữ liệu đa đơn vị",
        "Phụ huynh theo dõi; Nhà trường/trung tâm chịu trách nhiệm cập nhật dữ liệu nguồn.",
        [
            ("Tổng quan", "Số con, lớp, lịch và việc cần chú ý.", "Số học viên/lớp/khóa và lịch sắp tới."),
            ("Hồ sơ con", "Thông tin cơ bản, trường, lớp và trạng thái.", "Thông tin học viên, trung tâm, lớp/khóa."),
            ("Học tập", "Điểm danh, nhận xét phát triển, báo giảng.", "Chuyên cần, điểm/xếp loại, bài tập, chứng chỉ."),
            ("Học phí", "Khoản phải thu, đã thu, còn lại, lịch sử.", "Công nợ khóa, giáo trình, lệ phí, phiếu thu."),
            ("Xin phép", "Gửi đơn nghỉ theo con/lớp và theo dõi xử lý.", "Gửi đơn theo học viên/lớp/ca."),
            ("Trao đổi", "Xem lịch sử giáo viên/học vụ trao đổi.", "Xem phản hồi về chuyên cần/kết quả/lộ trình."),
            ("Thông báo", "Nhận theo trường/lớp/cá nhân.", "Nhận theo trung tâm/lớp/cá nhân."),
            ("Đa đơn vị", "Nhóm dữ liệu theo từng trường/trung tâm.", "Không trộn lớp, lịch, công nợ và thông báo."),
        ],
        [
            ("Đăng nhập bằng tài khoản phụ huynh.", "Tổng quan hiển thị con học mầm non."),
            ("Mở Hồ sơ con, Học tập và Lịch.", "Dữ liệu thuộc đúng trường/lớp."),
            ("Mở Học phí.", "Đối chiếu công nợ và phiếu thu."),
            ("Gửi đơn Xin phép.", "Chọn đúng con, lớp và thời gian."),
            ("Mở Trao đổi và Thông báo.", "Nhìn thấy đúng nguồn trường."),
        ],
        [
            ("Từ cùng tài khoản, chọn học viên tại trung tâm.", "Dữ liệu chuyển nhóm trung tâm."),
            ("Mở lịch, kết quả và bài tập.", "Đúng lớp/khóa ngoại ngữ."),
            ("Mở học phí.", "Không trộn công nợ mầm non."),
            ("Gửi đơn xin nghỉ theo ca.", "Đơn đến đúng trung tâm."),
            ("Mở thông báo.", "Nhận đủ nguồn nhưng mỗi bản ghi ghi rõ đơn vị."),
        ],
        [
            ("TC-4.11-01", "Một phụ huynh có hai con ở hai đơn vị.", "Thấy đủ hai nhóm dữ liệu."),
            ("TC-4.11-02", "Mở hồ sơ học sinh không liên kết.", "Bị chặn."),
            ("TC-4.11-03", "Hai đơn vị gửi Toàn đơn vị.", "Nhận hai thông báo, không mất bản ghi."),
            ("TC-4.11-04", "Gửi đơn cho lớp đã kết thúc.", "Không cho chọn hoặc báo không hợp lệ."),
        ],
    )

    module_guide(
        doc,
        12,
        "Báo cáo, cài đặt và chức năng dùng chung",
        "Quản trị/Quản lý/Kế toán theo quyền; mọi người dùng quản lý hồ sơ cá nhân.",
        [
            ("Báo cáo tài chính", "Lọc theo ngày; tổng thu, hoàn phí, công nợ, chi, lãi/lỗ.", "Áp dụng cùng chỉ số cho trung tâm."),
            ("Xuất dữ liệu", "Xuất CSV thu theo kỳ.", "Xuất CSV đúng trung tâm và khoảng ngày."),
            ("Cài đặt đăng nhập", "Số lần sai, thời gian khóa, độ dài mật khẩu.", "Cấu hình hệ thống theo phạm vi quyền."),
            ("Mẫu in", "Logo, chữ ký và footer của trường.", "Logo, chữ ký và footer của trung tâm."),
            ("Thông tin cá nhân", "Ảnh đại diện, họ tên, email, điện thoại.", "Áp dụng cho mọi vai trò."),
            ("Chuyển đơn vị", "Chỉ hiện khi tài khoản được gán nhiều đơn vị.", "Tải lại toàn bộ dữ liệu theo đơn vị mới."),
            ("Cảnh báo chưa lưu", "Xác nhận trước khi rời form đang sửa.", "Áp dụng toàn ứng dụng."),
        ],
        [
            ("Mở Báo cáo tài chính và chọn khoảng ngày.", "Đối chiếu tổng thu/chi với chứng từ trường."),
            ("Chọn Xuất CSV.", "Mở tệp và kiểm tra các kỳ thu."),
            ("Mở Cài đặt nếu có quyền.", "Kiểm tra chính sách đăng nhập và mẫu in."),
            ("Mở Thông tin cá nhân.", "Cập nhật hồ sơ và lưu."),
            ("Thử rời một form đang sửa.", "Hệ thống cảnh báo chưa lưu."),
        ],
        [
            ("Chuyển sang Trung tâm Ngoại ngữ.", "Dashboard và báo cáo tải lại."),
            ("Lọc báo cáo theo thời gian khóa học.", "Số liệu không chứa trường mầm non."),
            ("Xuất CSV.", "Tên kỳ/khóa và số tiền đúng."),
            ("Kiểm tra mẫu in trung tâm.", "Logo/nhãn ký/footer đúng đơn vị."),
            ("Cập nhật hồ sơ cá nhân.", "Không làm thay đổi hồ sơ người dùng khác."),
        ],
        [
            ("TC-4.12-01", "Kế toán xem báo cáo.", "Được phép; không sửa lớp/học sinh ngoài quyền."),
            ("TC-4.12-02", "Người không có quyền mở Cài đặt.", "Hiển thị không có quyền."),
            ("TC-4.12-03", "Chuyển đơn vị.", "Không còn dữ liệu cache của đơn vị cũ."),
            ("TC-4.12-04", "Rời form chưa lưu.", "Có lựa chọn Ở lại/Rời trang."),
        ],
    )

    doc.add_heading("4.13. Danh mục chức năng nhỏ không được bỏ qua", level=2)
    base.add_table(
        doc,
        ["Nhóm", "Chức năng cần kiểm tra"],
        [
            ("Đăng nhập", "Ghi nhớ tên đăng nhập; sai mật khẩu; khóa tài khoản; đổi mật khẩu lần đầu; đăng xuất."),
            ("Danh sách", "Tìm kiếm; bộ lọc; số lượng bản ghi; trạng thái rỗng; mở chi tiết; quay lại danh sách."),
            ("Biểu mẫu", "Trường bắt buộc; định dạng ngày/số tiền/email/điện thoại; nút Lưu; cảnh báo chưa lưu; thông báo thành công/lỗi."),
            ("Trạng thái", "Kích hoạt/ngừng; mở/đóng kỳ; bắt đầu/kết thúc/hủy buổi; tiếp nhận/đang học/bảo lưu/ngừng/hoàn thành."),
            ("Quan hệ", "Gán/gỡ vai trò; liên kết/gỡ phụ huynh; phân công/gỡ giáo viên; xếp/chuyển/ngừng/hoàn thành lớp."),
            ("Chứng từ", "Sinh số; mở chi tiết; in; in lại; tách dữ liệu theo đơn vị; không sửa bản ghi cũ."),
            ("Thông báo", "Phạm vi; lớp/cá nhân bắt buộc; đính kèm; đơn vị gửi; hiển thị đa đơn vị."),
            ("Bảo mật", "Chặn URL ngoài quyền; chặn dữ liệu ngoài đơn vị; nhật ký người thao tác; không lộ mật khẩu sau lần tạo."),
            ("Khả dụng", "Menu đúng vai trò; tiêu đề trang đúng; trạng thái tải/lỗi/rỗng; hiển thị trên kích thước màn hình phổ biến."),
        ],
        [1.3, 5.2],
    )

    doc.add_heading("5. Tài khoản và dữ liệu demo", level=1)
    base.add_table(
        doc,
        ["Vai trò", "Mầm non", "Ngoại ngữ", "Mật khẩu"],
        [
            ("Quản lý", "demo_quanly_mn", "demo_quanly_nn", "Edu@123Qaz"),
            ("Tuyển sinh", "demo_tuyensinh_mn", "demo_tuyensinh_nn", "Edu@123Qaz"),
            ("Học vụ", "demo_hocvu_mn", "demo_hocvu_nn", "Edu@123Qaz"),
            ("Giáo viên", "demo_giaovien_mn", "demo_giaovien_nn", "Edu@123Qaz"),
            ("Kế toán", "demo_ketoan_mn", "demo_ketoan_nn", "Edu@123Qaz"),
            ("Phụ huynh", "0988002026", "Dùng chung", "Edu@123Qaz"),
        ],
        [1.25, 1.8, 1.8, 1.65],
    )
    base.add_callout(
        doc,
        "Bảo mật:",
        "Các tài khoản và mật khẩu trên chỉ dùng cho môi trường demo/đào tạo.",
        base.AMBER,
    )

    doc.add_heading("6. Checklist nghiệm thu cuối", level=1)
    base.add_table(
        doc,
        ["Mã", "Nội dung", "Tiêu chí đạt"],
        [
            ("UAT-01", "Tài khoản nhân sự", "Đủ vai trò, đúng đơn vị; giáo viên tạo từ hồ sơ."),
            ("UAT-01B", "Tài khoản giáo viên", "Tên theo gv_ten.ho; trùng thì thêm số; không dùng điện thoại."),
            ("UAT-02", "Tuyển sinh", "Lead có lịch sử và chuyển thành học sinh Tiếp nhận."),
            ("UAT-03", "Đào tạo", "Có chương trình, lớp, giáo viên, học sinh và lịch."),
            ("UAT-04", "Vận hành", "Điểm danh, báo giảng, xin phép và trao đổi hoạt động."),
            ("UAT-05", "Tài chính", "Có kỳ thu, công nợ, phiếu thu và báo cáo khớp."),
            ("UAT-05B", "Phiếu in", "Phiếu nhập học, xếp lớp và phiếu thu mở/in lại đúng dữ liệu."),
            ("UAT-06", "Kết quả", "Đánh giá/chứng chỉ hiển thị đúng Portal."),
            ("UAT-07", "Phân quyền", "Không vai trò nào truy cập dữ liệu ngoài phạm vi."),
        ],
        [0.9, 2.4, 3.2],
    )

    doc.core_properties.title = "Hướng dẫn vận hành tuần tự QLTruongHoc"
    doc.core_properties.author = "Nhóm phát triển QLTruongHoc"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
