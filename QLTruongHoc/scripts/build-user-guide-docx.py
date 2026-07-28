from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "HUONG_DAN_SU_DUNG_VA_KIEM_THU_QLTRUONGHOC.docx"
IMG = ROOT / "docs" / "huong-dan-su-dung" / "images"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F8FC"
GREEN = "E5F4EC"
AMBER = "FFF3D6"
GRAY = "667085"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mục lục sẽ được cập nhật khi mở tài liệu trong Microsoft Word."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2, placeholder, fld_char3])


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
        if widths:
            cell.width = Inches(widths[idx])
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
            if widths:
                cell.width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    run = p.add_run(title + " ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    p.add_run(text)
    doc.add_paragraph()


def add_steps(doc, steps):
    rows = []
    for idx, (action, expected) in enumerate(steps, 1):
        rows.append((idx, action, expected))
    return add_table(doc, ["Bước", "Thao tác", "Kết quả cần đạt"], rows, [0.55, 3.25, 2.7])


def add_test_cases(doc, rows):
    return add_table(
        doc,
        ["Mã", "Điều kiện / thao tác kiểm thử", "Kết quả mong đợi"],
        rows,
        [0.85, 3.15, 2.5],
    )


def add_module_sop(
    doc,
    code,
    title,
    business_case,
    roles,
    prerequisites,
    steps,
    tests,
    differences=None,
):
    doc.add_heading(f"{code} — {title}", level=2)
    doc.add_paragraph(f"Business case: {business_case}")
    add_callout(
        doc,
        "Vai trò thực hiện:",
        roles + "  Điều kiện trước: " + prerequisites,
    )
    doc.add_heading("Các bước thực hiện", level=3)
    add_steps(doc, steps)
    if differences:
        doc.add_heading("Khác biệt theo loại đơn vị", level=3)
        add_table(
            doc,
            ["Loại đơn vị", "Cách áp dụng"],
            differences,
            [1.7, 4.8],
        )
    doc.add_heading("Test case và tiêu chí nghiệm thu", level=3)
    add_test_cases(doc, tests)


def add_figure(doc, filename, caption, width=6.35):
    path = IMG / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(34, 40, 49)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 26, DEEP_BLUE, 0, 10),
        ("Subtitle", 13, GRAY, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_after = Pt(8)

    if "Guide Bullet" not in styles:
        bullet = styles.add_style("Guide Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = styles["Guide Bullet"]
    bullet.base_style = styles["Normal"]
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "QLTRUONGHOC  |  HƯỚNG DẪN SỬ DỤNG VÀ KIỂM THỬ"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    core = doc.core_properties
    core.title = "Hướng dẫn sử dụng và kiểm thử hệ thống QLTruongHoc"
    core.author = "Nhóm phát triển QLTruongHoc"
    core.subject = "Hướng dẫn tạo dữ liệu, sử dụng và kiểm thử nghiệp vụ"
    core.keywords = "QLTruongHoc, hướng dẫn sử dụng, kiểm thử, đào tạo"


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("QLTRUONGHOC")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HƯỚNG DẪN SỬ DỤNG\nVÀ KIỂM THỬ HỆ THỐNG")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tạo dữ liệu mẫu · Thực hành nghiệp vụ · Xác nhận kết quả")

    doc.add_paragraph()
    line = doc.add_table(rows=1, cols=1)
    line.style = "Table Grid"
    set_repeat_table_header(line.rows[0])
    cell = line.cell(0, 0)
    set_cell_shading(cell, BLUE)
    cell.paragraphs[0].add_run(" ")

    for _ in range(4):
        doc.add_paragraph()
    info = add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Tác giả", "Nhóm phát triển QLTruongHoc"),
            ("Phiên bản", "2.0"),
            ("Ngày phát hành", "28/07/2026"),
            ("Phạm vi", "Trường Mầm non Hoa Nắng và Trung tâm Ngoại ngữ Quận 8"),
        ],
        [1.6, 4.9],
    )
    info.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("Kiểm soát tài liệu", level=1)
    add_table(
        doc,
        ["Phiên bản", "Ngày", "Nội dung", "Người thực hiện"],
        [
            ("1.0", "28/07/2026", "Phát hành hướng dẫn tạo dữ liệu và kiểm thử các luồng chính.", "Nhóm phát triển QLTruongHoc"),
            ("2.0", "28/07/2026", "Bổ sung SOP chi tiết cho toàn bộ vòng đời vận hành theo từng module và loại đơn vị.", "Nhóm phát triển QLTruongHoc"),
        ],
        [0.9, 1.1, 3.5, 1.0],
    )
    add_callout(
        doc,
        "Đối tượng sử dụng:",
        "Quản trị hệ thống, quản lý đơn vị, tuyển sinh/tư vấn, học vụ, giáo viên, kế toán, phụ huynh và nhân sự kiểm thử.",
    )

    doc.add_heading("Mục lục", level=1)
    add_toc_field(doc.add_paragraph())
    doc.add_paragraph(
        "Nếu mục lục chưa hiển thị số trang, mở tài liệu bằng Microsoft Word, nhấp vào mục lục và chọn “Update Table”."
    )
    doc.add_page_break()

    doc.add_heading("1. Mục đích và phạm vi", level=1)
    doc.add_paragraph(
        "Tài liệu hướng dẫn cách chuẩn bị dữ liệu mẫu và kiểm tra các luồng nghiệp vụ chính của hệ thống quản lý trường học/trung tâm đào tạo. "
        "Các bước được sắp xếp theo quan hệ phụ thuộc dữ liệu: tài khoản → chương trình/lớp → học sinh → lịch học → vận hành → tài chính → portal."
    )
    add_callout(
        doc,
        "Nguyên tắc thực hiện:",
        "Hoàn thành một business case và chạy các test case đi kèm trước khi chuyển sang business case tiếp theo.",
        GREEN,
    )
    add_table(
        doc,
        ["Nhóm", "Phạm vi"],
        [
            ("Đơn vị mầm non", "Chương trình 5–6 tuổi, lớp mầm non, lịch cả ngày, đánh giá theo lĩnh vực phát triển."),
            ("Trung tâm ngoại ngữ", "Chương trình A2, lớp theo ca, kết quả có điểm số."),
            ("Portal", "Giáo viên, kế toán và phụ huynh có hai con học tại hai đơn vị."),
            ("Ngoài phạm vi", "Kiểm thử tải, sao lưu/khôi phục và cấu hình quyền hạt chi tiết."),
        ],
        [1.7, 4.8],
    )

    doc.add_heading("2. Chuẩn bị môi trường và dữ liệu mẫu", level=1)
    doc.add_heading("2.1 Điều kiện trước khi thực hiện", level=2)
    for item in (
        "Ứng dụng và cơ sở dữ liệu đã được cấu hình; truy cập giao diện tại http://localhost:5173.",
        "Đã có các đơn vị SYSTEM, MN-HOA-NANG và TTNN-Q8 cùng tám vai trò hệ thống.",
        "Người thực hiện có quyền chạy lệnh tại thư mục dự án.",
    ):
        doc.add_paragraph("• " + item, style="Guide Bullet")

    doc.add_heading("2.2 Tạo dữ liệu tự động", level=2)
    doc.add_paragraph("Tại thư mục dự án, chạy lần lượt:")
    for command in ("pnpm db:seed:test:mam-non", "pnpm db:seed:test:ngoai-ngu"):
        p = doc.add_paragraph()
        p.style = "Intense Quote"
        p.add_run(command)
    add_callout(
        doc,
        "Có thể chạy lại an toàn:",
        "Bộ seed đối chiếu theo dữ liệu nghiệp vụ đã có, bổ sung phần còn thiếu và không tạo trùng lớp, học sinh, kỳ thu, đánh giá hoặc thông báo.",
    )

    doc.add_heading("2.3 Bộ tài khoản thực hành", level=2)
    add_table(
        doc,
        ["Vai trò", "Mầm non", "Ngoại ngữ", "Mật khẩu"],
        [
            ("Quản lý đơn vị", "demo_quanly_mn", "demo_quanly_nn", "Edu@123Qaz"),
            ("Học vụ", "demo_hocvu_mn", "demo_hocvu_nn", "Edu@123Qaz"),
            ("Giáo viên", "demo_giaovien_mn", "demo_giaovien_nn", "Edu@123Qaz"),
            ("Kế toán", "demo_ketoan_mn", "demo_ketoan_nn", "Edu@123Qaz"),
            ("Phụ huynh dùng chung", "0988002026", "0988002026", "Edu@123Qaz"),
        ],
        [1.45, 1.75, 1.75, 1.55],
    )
    add_callout(
        doc,
        "Bảo mật:",
        "Các tài khoản trên chỉ dùng cho môi trường kiểm thử/đào tạo. Không sử dụng mật khẩu mẫu trên môi trường thật.",
        AMBER,
    )

    doc.add_heading("2.4 Kết quả dữ liệu sau khi tạo", level=2)
    add_table(
        doc,
        ["Đơn vị", "Lớp / học viên", "Lịch học", "Tài chính"],
        [
            ("Mầm non Hoa Nắng", "1 lớp · 5 trẻ", "21 buổi từ 28/07 đến 25/08/2026", "1 kỳ thu · 5 khoản phải thu"),
            ("Ngoại ngữ Quận 8", "1 lớp · 4 học viên", "12 buổi từ 28/07 đến 25/08/2026", "1 kỳ thu · 4 khoản phải thu"),
        ],
        [1.65, 1.55, 2.05, 1.25],
    )

    doc.add_heading("3. Quy trình vận hành toàn trình theo module", level=1)
    doc.add_paragraph(
        "Các module dưới đây được sắp xếp theo thứ tự dữ liệu phát sinh trong thực tế. "
        "Không nên đảo thứ tự vì module sau sử dụng dữ liệu do module trước tạo ra."
    )

    add_module_sop(
        doc,
        "M01",
        "Khởi tạo đơn vị và tài khoản nhân sự",
        "Mỗi nhân sự phải có tài khoản, vai trò và đơn vị làm việc chính xác trước khi phát sinh nghiệp vụ.",
        "Quản trị hệ thống hoặc quản lý đơn vị có quyền quản lý người dùng.",
        "Đơn vị và danh mục vai trò đã tồn tại.",
        [
            ("Mở Hệ thống → Người dùng → Tạo người dùng.", "Hiển thị biểu mẫu tạo tài khoản."),
            ("Tạo tài khoản Quản lý đơn vị; nhập họ tên, tên đăng nhập, email/điện thoại.", "Tài khoản được gắn vai trò quan_ly_don_vi tại đúng đơn vị."),
            ("Tạo lần lượt tài khoản Kế toán, Học vụ, Tuyển sinh hoặc Tư vấn.", "Mỗi tài khoản có đúng một vai trò nghiệp vụ khởi tạo."),
            ("Đăng nhập thử từng tài khoản.", "Portal và menu thay đổi đúng vai trò."),
            ("Không tạo Giáo viên hoặc Phụ huynh tại màn hình này.", "Hai nhóm này được tạo từ hồ sơ nghiệp vụ tương ứng."),
        ],
        [
            ("TC-M01-01", "Tạo học vụ tại đơn vị mầm non.", "Thấy Đào tạo, Học sinh, Lớp, Lịch; không có quyền thu tiền."),
            ("TC-M01-02", "Tạo kế toán tại trung tâm ngoại ngữ.", "Thấy Tài chính; không sửa lớp hoặc điểm danh."),
            ("TC-M01-03", "Tạo trùng tên đăng nhập.", "Hệ thống báo trùng và không tạo bản ghi mới."),
            ("TC-M01-04", "Gán học vụ/giáo viên tại node hệ thống.", "Bị chặn vì node hệ thống không vận hành lớp."),
        ],
        [
            ("Đơn vị hệ thống", "Chỉ tạo quản trị, quản lý được ủy quyền hoặc kế toán tổng hợp; không tạo giáo viên/học vụ/tuyển sinh."),
            ("Trường mầm non", "Tạo quản lý, học vụ, kế toán, tuyển sinh; giáo viên tạo từ hồ sơ Giáo viên."),
            ("Trung tâm ngoại ngữ", "Tương tự mầm non; có thể dùng vai trò tư vấn cho quy trình chăm sóc lead."),
        ],
    )

    add_module_sop(
        doc,
        "M02",
        "Tạo hồ sơ giáo viên và tài khoản đăng nhập",
        "Hồ sơ giáo viên là dữ liệu nhân sự giảng dạy; tài khoản đăng nhập phải liên kết với đúng hồ sơ để giới hạn lớp và lịch được phân công.",
        "Học vụ hoặc quản lý đơn vị.",
        "Đơn vị đã có tài khoản học vụ/quản lý.",
        [
            ("Mở Đào tạo → Giáo viên → Tạo giáo viên.", "Hiển thị biểu mẫu hồ sơ giáo viên."),
            ("Nhập họ tên, điện thoại, email, chuyên môn và trình độ.", "Hồ sơ giáo viên được tạo với mã tự sinh."),
            ("Mở chi tiết giáo viên → Tạo tài khoản đăng nhập.", "Hiển thị tên đăng nhập và mật khẩu tạm."),
            ("Xác nhận tạo tài khoản.", "Tài khoản mang vai trò giao_vien và liên kết với hồ sơ."),
            ("Đăng nhập bằng tài khoản mới.", "Portal giáo viên chỉ hiển thị lớp được phân công."),
        ],
        [
            ("TC-M02-01", "Tạo tài khoản từ hồ sơ chưa có tài khoản.", "Tạo thành công và gắn nguoiDungId."),
            ("TC-M02-02", "Tạo tài khoản lần hai cho cùng giáo viên.", "Hệ thống báo giáo viên đã có tài khoản."),
            ("TC-M02-03", "Giáo viên chưa được phân công lớp đăng nhập.", "Portal hiển thị 0 lớp, không lộ lớp khác."),
        ],
        [
            ("Mầm non", "Chuyên môn chăm sóc/giáo dục theo độ tuổi; thường phân công vai trò Chủ nhiệm."),
            ("Ngoại ngữ", "Chuyên môn theo ngôn ngữ/trình độ; thường phân công Giáo viên chính hoặc Trợ giảng."),
        ],
    )

    add_module_sop(
        doc,
        "M03",
        "Tiếp nhận và chăm sóc tuyển sinh",
        "Quản lý người quan tâm từ lúc tiếp nhận đến khi đủ điều kiện đăng ký, bảo toàn lịch sử tư vấn và nguồn tuyển sinh.",
        "Tuyển sinh hoặc Tư vấn.",
        "Đã có tài khoản tuyển sinh/tư vấn tại đơn vị.",
        [
            ("Mở Tuyển sinh → Tạo lead.", "Biểu mẫu tiếp nhận người quan tâm được mở."),
            ("Nhập người liên hệ, số điện thoại, nguồn, độ tuổi/trình độ và nhu cầu.", "Lead được tạo ở trạng thái Mới với mã tự sinh."),
            ("Mở chi tiết lead → Ghi nhận hoạt động.", "Lưu cuộc gọi, tin nhắn, gặp trực tiếp, hẹn lịch hoặc học thử."),
            ("Cập nhật trạng thái phù hợp sau hoạt động.", "Lead chuyển Mới → Đang chăm sóc → Đã hẹn lịch/Đã học thử."),
            ("Nếu không tiếp tục, nhập lý do.", "Lead đóng nhưng vẫn giữ toàn bộ lịch sử."),
        ],
        [
            ("TC-M03-01", "Tạo lead đủ họ tên và số điện thoại.", "Mã LD theo năm được sinh, trạng thái Mới."),
            ("TC-M03-02", "Ghi hoạt động và đổi trạng thái.", "Hoạt động append-only và trạng thái cập nhật đồng thời."),
            ("TC-M03-03", "Đánh dấu Không tiếp tục nhưng bỏ trống lý do.", "Bị chặn."),
            ("TC-M03-04", "Xem lead từ đơn vị khác.", "Không hiển thị do cách ly dữ liệu."),
        ],
        [
            ("Mầm non", "Trường Độ tuổi/trình độ ghi độ tuổi của trẻ và nhu cầu bán trú/chăm sóc."),
            ("Ngoại ngữ", "Ghi trình độ hiện tại, mục tiêu, ca học; có thể dùng hoạt động Học thử."),
        ],
    )

    add_module_sop(
        doc,
        "M04",
        "Xác nhận đăng ký và hoàn thiện hồ sơ học sinh",
        "Khi lead đồng ý đăng ký, hệ thống tạo hồ sơ học sinh và phụ huynh chính thức nhưng chưa mặc định coi học sinh đã vào lớp.",
        "Tuyển sinh xác nhận; Học vụ hoàn thiện hồ sơ.",
        "Lead chưa ở trạng thái Đã đăng ký hoặc Không tiếp tục.",
        [
            ("Từ chi tiết lead chọn Xác nhận đăng ký.", "Mở biểu mẫu thông tin học sinh."),
            ("Nhập họ tên học sinh, ngày sinh, giới tính, địa chỉ và ngày nhập học.", "Tạo hồ sơ học sinh trạng thái Tiếp nhận."),
            ("Chọn quan hệ của người liên hệ với học sinh.", "Tạo hoặc tái sử dụng hồ sơ phụ huynh theo số điện thoại."),
            ("Mở Học sinh → chi tiết hồ sơ.", "Kiểm tra mã học sinh, thông tin y tế, liên hệ khẩn cấp."),
            ("Bổ sung thông tin còn thiếu và lưu.", "Hồ sơ đầy đủ nhưng chưa có lớp đang học."),
            ("Từ liên kết phụ huynh chọn Tạo tài khoản đăng nhập.", "Phụ huynh nhận tài khoản tạm và xem đúng con."),
        ],
        [
            ("TC-M04-01", "Xác nhận lead hợp lệ.", "Tạo học sinh Tiếp nhận và liên kết phụ huynh chính."),
            ("TC-M04-02", "Xác nhận lại lead đã đăng ký.", "Bị chặn, không tạo học sinh trùng."),
            ("TC-M04-03", "Số điện thoại phụ huynh đã tồn tại.", "Tái sử dụng đúng hồ sơ; yêu cầu xác nhận nếu khác đơn vị."),
            ("TC-M04-04", "Sau xác nhận nhưng chưa xếp lớp.", "Portal không hiển thị học sinh là Đang học."),
        ],
    )

    add_module_sop(
        doc,
        "M05",
        "Thiết lập chương trình đào tạo và lớp học",
        "Chương trình mô tả khung đào tạo; lớp là đơn vị tổ chức học thực tế có thời gian, sĩ số và phòng học.",
        "Học vụ hoặc quản lý đơn vị.",
        "Đã có đơn vị và giáo viên.",
        [
            ("Mở Đào tạo → Chương trình → Tạo mới.", "Tạo tên chương trình, cấp độ, số buổi/giờ và mô tả."),
            ("Mở Lớp học → Tạo lớp.", "Chọn chương trình, nhập tên lớp, cấp độ, ngày bắt đầu/kết thúc, sĩ số và phòng."),
            ("Lưu lớp ở trạng thái Chuẩn bị.", "Lớp tồn tại nhưng chưa vận hành."),
            ("Mở chi tiết lớp → Phân công giáo viên.", "Chọn giáo viên, vai trò và ngày hiệu lực."),
            ("Khi đủ điều kiện chọn Bắt đầu/Đang học.", "Lớp sẵn sàng nhận học sinh và sinh lịch."),
        ],
        [
            ("TC-M05-01", "Tạo lớp thiếu chương trình.", "Bị chặn."),
            ("TC-M05-02", "Phân công giáo viên đúng đơn vị.", "Giáo viên thấy lớp trên Portal."),
            ("TC-M05-03", "Phân công giáo viên của đơn vị khác.", "Không cho chọn hoặc API từ chối."),
            ("TC-M05-04", "Ngày kết thúc trước ngày bắt đầu.", "Bị chặn bởi validation."),
        ],
        [
            ("Mầm non", "Chương trình theo độ tuổi; lịch thường cả ngày, Chủ nhiệm là vai trò trung tâm."),
            ("Ngoại ngữ", "Chương trình theo cấp độ/khóa; lớp theo ca và có thể có giáo viên chính/trợ giảng."),
        ],
    )

    add_module_sop(
        doc,
        "M06",
        "Xếp lớp, chuyển lớp và kết thúc học",
        "Học sinh chỉ trở thành Đang học khi có lượt xếp lớp hiệu lực; mọi thay đổi lớp phải lưu được lịch sử.",
        "Học vụ; quản lý đơn vị xử lý trường hợp vượt sĩ số.",
        "Học sinh Tiếp nhận, lớp Đang học và còn hiệu lực.",
        [
            ("Mở chi tiết lớp → Xếp học sinh.", "Danh sách chỉ gồm học sinh đủ điều kiện."),
            ("Chọn học sinh và ngày vào lớp.", "Tạo lượt xếp lớp Đang học."),
            ("Kiểm tra hồ sơ học sinh.", "Trạng thái tổng chuyển thành Đang học."),
            ("Khi chuyển lớp, chọn lớp đích và ngày chuyển.", "Lượt cũ đóng Chuyển lớp; lượt mới mở cùng ngày."),
            ("Khi nghỉ/hoàn thành, kết thúc lượt xếp lớp với lý do.", "Nếu không còn lượt active, trạng thái học sinh đồng bộ."),
        ],
        [
            ("TC-M06-01", "Xếp học sinh Tiếp nhận vào lớp.", "Tạo enrollment và đồng bộ Đang học."),
            ("TC-M06-02", "Xếp học sinh vào lớp đủ sĩ số.", "Người thường bị chặn; quản lý có quyền vượt sĩ số."),
            ("TC-M06-03", "Kết thúc lượt active cuối cùng.", "Học sinh chuyển Hoàn thành/Ngừng học."),
            ("TC-M06-04", "Chuyển lớp.", "Không tồn tại hai lượt Đang học chồng nhau trái quy tắc."),
        ],
    )

    add_module_sop(
        doc,
        "M07",
        "Lập thời khóa biểu và vận hành buổi học",
        "Quy tắc lịch lặp tạo các buổi học cụ thể; buổi học là đơn vị để bắt đầu học, điểm danh và báo giảng.",
        "Học vụ tạo lịch; Giáo viên vận hành buổi học.",
        "Lớp Đang học và đã phân công giáo viên.",
        [
            ("Mở Lịch học → Tạo quy tắc.", "Chọn lớp, thứ trong tuần, giờ, phòng, giáo viên và ngày hiệu lực."),
            ("Chọn Sinh buổi học đến ngày.", "Các buổi học được tạo trong phạm vi lớp."),
            ("Giáo viên mở buổi học hôm nay → Bắt đầu buổi học.", "Trạng thái buổi chuyển sang Đang học."),
            ("Mở Điểm danh và lưu trạng thái từng học sinh.", "Dữ liệu có mặt/vắng/đi trễ/về sớm được ghi nhận."),
            ("Mở Báo giảng và nhập nội dung, bài tập, ghi chú.", "Báo giảng gắn đúng buổi học."),
            ("Kết thúc buổi học.", "Buổi chuyển Đã học và số liệu Portal cập nhật."),
        ],
        [
            ("TC-M07-01", "Sinh lịch 28 ngày.", "Không tạo trùng buổi khi chạy lại."),
            ("TC-M07-02", "Điểm danh trước khi Bắt đầu buổi.", "Bị chặn."),
            ("TC-M07-03", "Điểm danh học sinh không thuộc roster ngày học.", "API từ chối."),
            ("TC-M07-04", "Ghi báo giảng cho buổi nghỉ/hủy.", "Bị chặn."),
        ],
        [
            ("Mầm non", "Thời lượng dài; điểm danh phục vụ đón/trả và theo dõi vắng."),
            ("Ngoại ngữ", "Ca ngắn; báo giảng nên ghi bài học, kỹ năng và bài tập về nhà."),
        ],
    )

    add_module_sop(
        doc,
        "M08",
        "Xin phép, cảnh báo vắng và trao đổi phụ huynh",
        "Nhà trường và gia đình cần một luồng có trạng thái rõ ràng cho xin nghỉ, phản hồi và lịch sử trao đổi.",
        "Phụ huynh gửi đơn; Giáo viên/Học vụ xử lý và trao đổi.",
        "Học sinh có lớp active và tài khoản phụ huynh.",
        [
            ("Phụ huynh mở tab Xin phép → Gửi đơn.", "Chọn lớp, thời gian nghỉ và lý do."),
            ("Học vụ mở Đơn xin phép.", "Danh sách đơn Chờ duyệt hiển thị đúng đơn vị/lớp."),
            ("Chọn Duyệt hoặc Từ chối, nhập ghi chú.", "Portal phụ huynh cập nhật trạng thái và nội dung phản hồi."),
            ("Giáo viên mở Trao đổi phụ huynh → Ghi nhận mới.", "Chọn học sinh, kênh, nội dung và kết quả."),
            ("Theo dõi cảnh báo vắng không phép.", "Portal hiển thị cảnh báo khi đạt ngưỡng."),
        ],
        [
            ("TC-M08-01", "Xin phép cho lớp không còn active.", "Lớp không xuất hiện trong danh sách."),
            ("TC-M08-02", "Từ chối đơn kèm ghi chú.", "Phụ huynh thấy Từ chối và ghi chú."),
            ("TC-M08-03", "Trao đổi bởi giáo viên.", "Portal hiển thị nhãn Giáo viên, không hiển thị enum."),
            ("TC-M08-04", "Vắng không phép 3 lần trong 7 ngày.", "Hiển thị cảnh báo cho đúng phụ huynh."),
        ],
    )

    add_module_sop(
        doc,
        "M09",
        "Thông báo và sự kiện",
        "Thông báo phải đến đúng người theo phạm vi toàn đơn vị, theo lớp hoặc cá nhân; phụ huynh nhiều đơn vị phải xem đủ nguồn gửi.",
        "Quản lý, học vụ hoặc nhân sự có quyền quản lý thông báo.",
        "Có đơn vị; với phạm vi lớp/cá nhân phải có lớp/học sinh.",
        [
            ("Mở Thông báo → Tạo mới.", "Nhập tiêu đề, nội dung và tệp/link nếu có."),
            ("Chọn phạm vi Toàn trường, Theo lớp hoặc Cá nhân.", "Biểu mẫu yêu cầu đúng đối tượng theo phạm vi."),
            ("Lưu thông báo.", "Thông báo có mã tự sinh và ghi đơn vị gửi."),
            ("Phụ huynh mở menu Thông báo.", "Chỉ nhận thông báo phù hợp liên kết con/lớp."),
            ("Chọn Xác nhận đã đọc.", "Trạng thái đọc cập nhật theo tài khoản."),
        ],
        [
            ("TC-M09-01", "Hai đơn vị cùng gửi Toàn trường.", "Phụ huynh có con ở hai đơn vị nhận đủ hai thông báo."),
            ("TC-M09-02", "Gửi Theo lớp.", "Chỉ phụ huynh có con active trong lớp nhận."),
            ("TC-M09-03", "Gửi Cá nhân.", "Chỉ phụ huynh của học sinh được chọn nhận."),
            ("TC-M09-04", "Thiếu đối tượng khi chọn Theo lớp/Cá nhân.", "Bị chặn."),
        ],
    )

    add_module_sop(
        doc,
        "M10",
        "Kỳ thu, thu tiền, điều chỉnh và chi phí",
        "Tài chính phải đi theo chuỗi danh mục khoản thu → kỳ thu → khoản phải thu → phiếu thu; chi phí và hoàn/chuyển phí cần tách người đề xuất và người duyệt.",
        "Kế toán lập nghiệp vụ; Quản lý đơn vị duyệt điều chỉnh/chi phí.",
        "Có lớp và học sinh đang học; đã có tài khoản kế toán.",
        [
            ("Tạo Danh mục khoản thu.", "Khai báo tên, loại, số tiền mặc định và tính bắt buộc."),
            ("Tạo Kỳ thu, chọn loại kỳ và thời gian.", "Kỳ ở trạng thái Nháp."),
            ("Áp dụng các khoản thu và Mở kỳ.", "Kỳ sẵn sàng sinh công nợ."),
            ("Sinh khoản phải thu theo lớp.", "Mỗi học sinh active nhận một khoản phải thu."),
            ("Mở khoản phải thu → Ghi nhận thu tiền.", "Tạo phiếu thu, cập nhật Đã thu và Còn lại."),
            ("Tạo yêu cầu hoàn phí/chuyển phí/bảo lưu.", "Yêu cầu ở trạng thái Chờ duyệt."),
            ("Quản lý đơn vị phê duyệt.", "Số liệu chỉ tác động báo cáo sau khi duyệt."),
            ("Mở Chi phí → Đề xuất chi.", "Chi phí chờ duyệt; sau duyệt được tính vào báo cáo."),
        ],
        [
            ("TC-M10-01", "Thu nhỏ hơn tổng phải thu.", "Trạng thái Thu một phần."),
            ("TC-M10-02", "Thu đủ số còn lại.", "Trạng thái Đã thu đủ."),
            ("TC-M10-03", "Sinh công nợ lại cùng kỳ/lớp.", "Không tạo trùng."),
            ("TC-M10-04", "Chi phí chưa duyệt.", "Không tính vào tổng chi đã ghi nhận."),
            ("TC-M10-05", "Kế toán hệ thống mở chi tiết kỳ của đơn vị con.", "Chỉ xem tổng hợp; thao tác chi tiết tại đơn vị sở hữu."),
        ],
        [
            ("Mầm non", "Thường có học phí, tiền ăn và dịch vụ theo tháng."),
            ("Ngoại ngữ", "Thường thu theo khóa, tài liệu, thi thử hoặc lệ phí chứng chỉ."),
        ],
    )

    add_module_sop(
        doc,
        "M11",
        "Đánh giá học tập, kết quả thi và chứng chỉ",
        "Kết quả phải gắn đúng lượt học/lớp; thành tích hoặc chứng chỉ có thể độc lập với một lớp cụ thể và được phụ huynh xem lại.",
        "Giáo viên hoặc Học vụ nhập; Phụ huynh chỉ xem.",
        "Học sinh có lượt xếp lớp.",
        [
            ("Mở hồ sơ học sinh → Kết quả học tập.", "Chọn lượt xếp lớp và loại đánh giá."),
            ("Nhập ngày, điểm/xếp loại và nhận xét.", "Kết quả gắn đúng enrollment."),
            ("Với mầm non chọn lĩnh vực phát triển.", "Chọn Thể chất, Nhận thức, Ngôn ngữ, Tình cảm–kỹ năng xã hội hoặc Thẩm mỹ."),
            ("Với lớp ngoại ngữ nhập điểm kiểm tra/thi và nhận xét kỹ năng.", "Điểm hiển thị đúng thang điểm đã dùng."),
            ("Mở Thành tích/Chứng chỉ → Thêm mới.", "Nhập tên, kết quả, ngày đạt, nơi cấp và minh chứng."),
            ("Phụ huynh mở tab Học tập.", "Hiển thị kết quả và chứng chỉ bằng nhãn tiếng Việt."),
        ],
        [
            ("TC-M11-01", "Đánh giá mầm non Theo tháng · Ngôn ngữ.", "Portal hiển thị đúng loại và lĩnh vực."),
            ("TC-M11-02", "Nhập lĩnh vực ngoài danh mục.", "API trả lỗi lĩnh vực không hợp lệ."),
            ("TC-M11-03", "Nhập điểm thi ngoại ngữ.", "Điểm và nhận xét gắn đúng lớp/lượt học."),
            ("TC-M11-04", "Thêm chứng chỉ có minh chứng.", "Chứng chỉ hiển thị trên hồ sơ và Portal."),
            ("TC-M11-05", "Xóa kết quả.", "Biến mất khỏi cả hồ sơ và Portal sau tải lại."),
        ],
        [
            ("Mầm non", "Đánh giá quá trình theo tháng/quý/năm và năm lĩnh vực phát triển; có thể không dùng điểm số."),
            ("Ngoại ngữ", "Ưu tiên điểm kiểm tra, xếp loại và nhận xét nghe–nói–đọc–viết; chứng chỉ/thi đầu ra."),
        ],
    )

    add_module_sop(
        doc,
        "M12",
        "Báo cáo, portal và kiểm soát quyền",
        "Portal tổng hợp dữ liệu từ các module; báo cáo và quyền truy cập phải phản ánh đúng đơn vị, vai trò và quan hệ dữ liệu.",
        "Tất cả vai trò trong phạm vi được cấp.",
        "Đã có dữ liệu từ M01 đến M11.",
        [
            ("Đăng nhập lần lượt Quản lý, Tuyển sinh, Học vụ, Giáo viên, Kế toán, Phụ huynh.", "Mỗi vai trò về đúng portal."),
            ("Đối chiếu chỉ số portal với danh sách chi tiết.", "Số lớp, lịch, công nợ, thông báo và kết quả khớp."),
            ("Đổi đơn vị với tài khoản đa đơn vị.", "Dữ liệu nghiệp vụ tải lại theo đơn vị đang chọn."),
            ("Mở URL ngoài quyền.", "Bị chặn ở cả giao diện và API."),
            ("Kiểm tra audit log cho thao tác quan trọng.", "Có người thực hiện, đơn vị, đối tượng và thời gian."),
        ],
        [
            ("TC-M12-01", "Giáo viên mở /finance.", "Bị điều hướng về trang hợp lệ."),
            ("TC-M12-02", "Kế toán đơn vị A xem đơn vị B.", "Bị chặn."),
            ("TC-M12-03", "Phụ huynh đổi tham số học sinh.", "API vẫn chốt theo guardian linkage."),
            ("TC-M12-04", "Quản trị hệ thống đăng nhập.", "Về cổng hệ thống, không bị thay bởi portal nhân viên đơn vị."),
        ],
    )

    doc.add_heading("4. Kiểm thử nhanh theo vai trò", level=1)
    doc.add_heading("BC-01 — Đăng nhập và nhận diện không gian làm việc", level=2)
    doc.add_paragraph(
        "Business case: Mỗi tài khoản phải vào đúng portal và đúng đơn vị theo vai trò. Menu chỉ hiển thị các chức năng được phép sử dụng."
    )
    add_steps(
        doc,
        [
            ("Mở địa chỉ ứng dụng.", "Hiển thị màn hình đăng nhập."),
            ("Nhập tên đăng nhập và mật khẩu của vai trò cần kiểm tra.", "Đăng nhập thành công."),
            ("Kiểm tra tên đơn vị ở thanh trên cùng và các nhóm menu bên trái.", "Đơn vị, vai trò và menu đúng với tài khoản."),
            ("Nhấn Đăng xuất sau khi hoàn thành.", "Phiên làm việc kết thúc và quay lại màn hình đăng nhập."),
        ],
    )
    add_figure(doc, "00-dang-nhap.png", "Hình 1. Màn hình đăng nhập hệ thống.")
    add_test_cases(
        doc,
        [
            ("TC-01", "Đăng nhập đúng thông tin.", "Mở đúng portal của vai trò."),
            ("TC-02", "Nhập sai mật khẩu.", "Hiển thị thông báo lỗi, không tạo phiên đăng nhập."),
            ("TC-03", "Mở trực tiếp URL ngoài quyền.", "Bị chặn hoặc chuyển về trang làm việc hợp lệ."),
        ],
    )

    doc.add_heading("5. Minh họa thiết lập dữ liệu đào tạo", level=1)
    doc.add_heading("BC-02 — Chương trình, giáo viên, lớp và lịch học", level=2)
    doc.add_paragraph(
        "Business case: Học vụ cần thiết lập khung đào tạo trước khi xếp học sinh. Lớp phải có chương trình, giáo viên phụ trách và lịch học mới có thể vận hành điểm danh/báo giảng."
    )
    add_steps(
        doc,
        [
            ("Đăng nhập tài khoản học vụ của đơn vị.", "Portal học vụ hiển thị đúng đơn vị."),
            ("Mở Chương trình đào tạo và kiểm tra chương trình mẫu.", "Có chương trình 5–6 tuổi hoặc A2."),
            ("Mở Giáo viên và kiểm tra hồ sơ giáo viên.", "Hồ sơ có chuyên môn, trình độ và tài khoản đăng nhập."),
            ("Mở Lớp học và chọn lớp mẫu.", "Lớp ở trạng thái Đang học, có giáo viên và học sinh."),
            ("Mở Lịch học, chọn khoảng ngày từ hiện tại.", "Các buổi học hiển thị theo quy tắc của đơn vị."),
        ],
    )
    add_figure(doc, "06-lich-hoc-tu-ngay-hien-tai.png", "Hình 2. Lịch học mầm non được sinh từ ngày hiện tại.")
    add_test_cases(
        doc,
        [
            ("TC-02.1", "Lọc lịch từ 28/07 đến 04/08/2026.", "Có 6 buổi lớp Lá Test Full."),
            ("TC-02.2", "Lọc theo giáo viên Test Mầm non.", "Chỉ còn lịch của giáo viên được phân công."),
            ("TC-02.3", "Kiểm tra lớp ngoại ngữ.", "Lịch theo ca 18:00–19:30 vào các ngày đã cấu hình."),
        ],
    )

    doc.add_heading("6. Minh họa vận hành của giáo viên", level=1)
    doc.add_heading("BC-03 — Portal giáo viên, điểm danh và báo giảng", level=2)
    doc.add_paragraph(
        "Business case: Giáo viên theo dõi lớp và lịch dạy của chính mình, vào nhanh các màn hình điểm danh, báo giảng và trao đổi phụ huynh."
    )
    add_steps(
        doc,
        [
            ("Đăng nhập demo_giaovien_mn.", "Mở Cổng giáo viên tại Trường Mầm non Hoa Nắng."),
            ("Kiểm tra bốn chỉ số đầu trang.", "1 lớp đang dạy; 1 buổi hôm nay; số báo giảng và trao đổi đúng dữ liệu."),
            ("Kiểm tra Lớp được phân công.", "Chỉ hiển thị Lá Test Full, vai trò chủ nhiệm."),
            ("Kiểm tra Lịch dạy 7 ngày tới.", "Có 6 buổi, đúng thời gian 07:30–16:30."),
            ("Mở Điểm danh hoặc Lớp học.", "Chỉ xem/thao tác dữ liệu thuộc lớp được phân công."),
        ],
    )
    add_figure(doc, "05-portal-giao-vien-mam-non.png", "Hình 3. Portal giáo viên mầm non với lớp và lịch dạy được phân công.")
    add_test_cases(
        doc,
        [
            ("TC-03.1", "Giáo viên mầm non mở Portal.", "Có đúng 1 lớp và 6 buổi trong 7 ngày."),
            ("TC-03.2", "Giáo viên nhập điểm danh khi buổi chưa bắt đầu.", "Hệ thống yêu cầu bắt đầu buổi học trước."),
            ("TC-03.3", "Giáo viên truy cập /finance.", "Bị chặn và không lộ dữ liệu tài chính."),
        ],
    )

    doc.add_heading("7. Minh họa nghiệp vụ tài chính", level=1)
    doc.add_heading("BC-04 — Kỳ thu, công nợ và phiếu thu", level=2)
    doc.add_paragraph(
        "Business case: Kế toán mở kỳ thu, sinh khoản phải thu theo lớp và ghi nhận thu tiền. Portal kế toán tổng hợp số liệu phục vụ đối chiếu hằng ngày."
    )
    add_steps(
        doc,
        [
            ("Đăng nhập demo_ketoan_mn.", "Mở Cổng kế toán của đơn vị mầm non."),
            ("Kiểm tra Kỳ thu đang mở.", "Có 1 kỳ thu tháng 07/2026."),
            ("Kiểm tra tổng công nợ và đã thu.", "Công nợ 16.300.000 đồng; đã thu 4.700.000 đồng."),
            ("Mở Học phí · Công nợ và chọn kỳ thu.", "Có 5 khoản phải thu."),
            ("Đối chiếu trạng thái.", "Có khoản Chưa thu, Thu một phần và Đã thu đủ."),
        ],
    )
    add_figure(doc, "07-portal-ke-toan.png", "Hình 4. Portal kế toán với kỳ thu, công nợ và số tiền đã thu.")
    add_test_cases(
        doc,
        [
            ("TC-04.1", "Ghi phiếu thu nhỏ hơn số còn lại.", "Khoản phải thu chuyển sang Thu một phần."),
            ("TC-04.2", "Ghi đủ số tiền còn lại.", "Khoản phải thu chuyển sang Đã thu đủ."),
            ("TC-04.3", "Kế toán đơn vị mở dữ liệu đơn vị khác.", "Bị chặn theo phạm vi đơn vị."),
        ],
    )

    doc.add_heading("8. Minh họa Portal phụ huynh đa đơn vị", level=1)
    doc.add_heading("BC-05 — Một phụ huynh, hai con, hai đơn vị", level=2)
    doc.add_paragraph(
        "Business case: Một phụ huynh có thể dùng một tài khoản để theo dõi nhiều con ở nhiều đơn vị. Dữ liệu phải được nhóm theo đơn vị và thông báo phải tổng hợp đúng phạm vi."
    )
    add_steps(
        doc,
        [
            ("Đăng nhập bằng 0988002026.", "Mở Cổng phụ huynh."),
            ("Kiểm tra các chỉ số Tổng quan.", "2 con, 2 lớp, 8 buổi sắp tới và 2 kết quả học tập."),
            ("Chọn Thông tin các con.", "Có nhóm Trung tâm Ngoại ngữ Quận 8 và Trường Mầm non Hoa Nắng."),
            ("Kiểm tra tab Học tập của từng con.", "Ngoại ngữ hiển thị điểm 8,5; mầm non hiển thị đánh giá Theo tháng · Ngôn ngữ."),
            ("Chọn menu Thông báo.", "Hiển thị 2 thông báo toàn đơn vị, mỗi thông báo ghi rõ đơn vị gửi."),
        ],
    )
    add_figure(doc, "01-portal-phu-huynh-tong-quan.png", "Hình 5. Tổng quan Portal phụ huynh với hai con và hai lớp.")
    add_figure(doc, "02-phu-huynh-hai-con-hai-don-vi.png", "Hình 6. Nhóm dữ liệu con tại Trung tâm Ngoại ngữ Quận 8.")
    add_figure(doc, "03-phu-huynh-con-mam-non-ket-qua.png", "Hình 7. Nhóm dữ liệu con tại Trường Mầm non Hoa Nắng.")
    add_figure(doc, "04-thong-bao-hai-don-vi.png", "Hình 8. Danh sách thông báo tổng hợp từ hai đơn vị.")
    add_test_cases(
        doc,
        [
            ("TC-05.1", "Đăng nhập phụ huynh dùng chung.", "Chỉ số Con đang theo dõi bằng 2."),
            ("TC-05.2", "Mở Thông tin các con.", "Có đúng hai nhóm đơn vị, không trộn lớp và kết quả."),
            ("TC-05.3", "Mở Thông báo.", "Có hai thông báo toàn đơn vị và đúng tên đơn vị gửi."),
            ("TC-05.4", "Phụ huynh A thử truy cập dữ liệu con của phụ huynh B.", "Bị chặn tại API; không hiển thị dữ liệu ngoài liên kết."),
        ],
    )

    doc.add_heading("9. Phân quyền theo vai trò và đơn vị", level=1)
    doc.add_heading("BC-06 — Xác nhận phạm vi làm việc", level=2)
    add_table(
        doc,
        ["Vai trò", "Tại đơn vị hệ thống", "Tại trường/trung tâm"],
        [
            ("Quản trị hệ thống", "Quản trị đơn vị, tài khoản, vai trò, nhật ký.", "Mở đơn vị để hỗ trợ quản trị; không thay thế nghiệp vụ hằng ngày."),
            ("Quản lý đơn vị", "Không áp dụng.", "Theo dõi tổng quan, duyệt tài chính/điều chỉnh và quản lý vận hành."),
            ("Tuyển sinh / tư vấn", "Xem tổng hợp khi được cấp.", "Quản lý lead, hoạt động chăm sóc và xác nhận đăng ký."),
            ("Học vụ", "Xem tổng hợp khi được cấp.", "Quản lý chương trình, lớp, lịch, học sinh và vận hành đào tạo."),
            ("Giáo viên", "Không áp dụng.", "Chỉ lớp được phân công; điểm danh, báo giảng, đánh giá, trao đổi."),
            ("Kế toán", "Xem báo cáo tổng hợp theo phân quyền.", "Quản lý kỳ thu, công nợ, phiếu thu, điều chỉnh và chi phí."),
            ("Phụ huynh", "Không thao tác nghiệp vụ.", "Xem dữ liệu các con, thông báo, học phí, xin phép và trao đổi."),
        ],
        [1.25, 2.45, 2.8],
    )
    add_test_cases(
        doc,
        [
            ("TC-06.1", "Đăng nhập lần lượt các tài khoản demo.", "Menu thay đổi đúng vai trò."),
            ("TC-06.2", "Đổi đơn vị (nếu tài khoản có nhiều đơn vị).", "Dữ liệu tải lại theo đơn vị đang chọn."),
            ("TC-06.3", "Gõ trực tiếp URL ngoài quyền.", "Bị chặn/điều hướng, không chỉ ẩn menu."),
        ],
    )

    doc.add_heading("10. Danh sách kiểm thử bàn giao", level=1)
    add_table(
        doc,
        ["Mã", "Nội dung", "Tiêu chí đạt"],
        [
            ("UAT-01", "Tạo lại dữ liệu mẫu", "Hai lệnh seed chạy thành công và chạy lặp không tăng số lượng."),
            ("UAT-02", "Portal giáo viên mầm non", "1 lớp, 1 buổi hôm nay, 6 buổi trong 7 ngày."),
            ("UAT-03", "Lịch học", "Dữ liệu bắt đầu từ ngày 28/07/2026 và đúng khung giờ."),
            ("UAT-04", "Portal kế toán", "1 kỳ mở; công nợ 16.300.000đ; đã thu 4.700.000đ."),
            ("UAT-05", "Portal phụ huynh", "2 con thuộc 2 đơn vị; kết quả học tập tách đúng."),
            ("UAT-06", "Thông báo phụ huynh", "Nhận đủ 2 thông báo toàn đơn vị."),
            ("UAT-07", "Phân quyền chéo", "Không vai trò nào truy cập được dữ liệu ngoài phạm vi."),
        ],
        [0.9, 2.65, 2.95],
    )
    add_callout(
        doc,
        "Cách ghi nhận kết quả:",
        "Đánh dấu Đạt/Không đạt cho từng UAT, đính kèm ảnh màn hình và ghi rõ tài khoản, đơn vị, thời điểm kiểm tra nếu có sai lệch.",
        GREEN,
    )

    doc.add_heading("11. Xử lý tình huống thường gặp", level=1)
    add_table(
        doc,
        ["Hiện tượng", "Cách kiểm tra và xử lý"],
        [
            ("Không thấy dữ liệu sau khi đăng nhập", "Kiểm tra đơn vị đang làm việc, tải lại trang và xác nhận seed của đúng đơn vị đã chạy."),
            ("Phụ huynh chỉ thấy một con", "Kiểm tra hai học sinh đã liên kết cùng hồ sơ phụ huynh theo số điện thoại 0988002026."),
            ("Phụ huynh không thấy thông báo", "Mở menu Thông báo; xác nhận liên kết phụ huynh bật Nhận thông báo và thông báo có phạm vi phù hợp."),
            ("Giáo viên không thấy lớp", "Kiểm tra hồ sơ giáo viên đã gắn tài khoản và phân công lớp còn hiệu lực."),
            ("Không có buổi để điểm danh", "Kiểm tra quy tắc lịch, khoảng ngày sinh buổi và trạng thái lớp."),
            ("Không sinh công nợ", "Kiểm tra kỳ thu đã mở, khoản thu đã áp dụng và học sinh đang có lượt xếp lớp hiệu lực."),
        ],
        [2.15, 4.35],
    )

    doc.add_heading("Phụ lục A — Dữ liệu tham chiếu", level=1)
    add_table(
        doc,
        ["Đối tượng", "Mầm non", "Ngoại ngữ"],
        [
            ("Lớp", "Lá Test Full · LOP0001", "English A2 Test Full · LOP0001"),
            ("Giáo viên", "Giáo viên Test Mầm non", "Giáo viên Test Tiếng Anh"),
            ("Học sinh đại diện", "Test Bé An", "Test Học viên Minh"),
            ("Phụ huynh dùng chung", "Nguyễn Minh Anh · 0988002026", "Nguyễn Minh Anh · 0988002026"),
            ("Kỳ thu", "Kỳ thu test mầm non 2026-07", "Kỳ thu test khóa tiếng Anh 2026-07"),
        ],
        [1.55, 2.5, 2.45],
    )

    doc.add_heading("Phụ lục B — Nhật ký thao tác UI theo từng bước", level=1)
    doc.add_paragraph(
        "Các ảnh dưới đây được chụp trực tiếp trên dữ liệu demo. Mỗi ảnh tương ứng "
        "một bước nghiệp vụ duy nhất; khi tạo nhiều bản ghi cùng loại, thực hiện lại "
        "đúng bước mẫu và thay dữ liệu đầu vào."
    )
    ui_steps = [
        ("B1", "Người thực hiện: Quản lý đơn vị. Kiểm tra tổng quan đơn vị trước khi tạo dữ liệu.", "M01-01-portal-quan-ly.png"),
        ("B2", "Người thực hiện: Quản lý đơn vị. Tạo tài khoản quản lý, tuyển sinh, tư vấn, học vụ hoặc kế toán.", "M01-02-tao-tai-khoan-nhan-su.png"),
        ("B3", "Người thực hiện: Học vụ/Quản lý. Mở biểu mẫu và tạo hồ sơ giáo viên.", "M02-01-tao-ho-so-giao-vien.png"),
        ("B4", "Người thực hiện: Học vụ/Quản lý. Xác nhận tài khoản đăng nhập đã liên kết đúng hồ sơ giáo viên.", "M02-02-tai-khoan-giao-vien-da-lien-ket.png"),
        ("B5", "Người thực hiện: Tuyển sinh/Tư vấn. Tạo và chăm sóc khách hàng tiềm năng.", "M03-01-tuyen-sinh.png"),
        ("B6", "Người thực hiện: Học vụ. Kiểm tra danh sách và mở hồ sơ học sinh/học viên.", "M04-01-ho-so-hoc-sinh.png"),
        ("B7", "Người thực hiện: Học vụ. Tạo lớp, gắn chương trình và theo dõi trạng thái lớp.", "M05-01-lop-hoc.png"),
        ("B8", "Người thực hiện: Học vụ/Quản lý. Phân công giáo viên, xếp học sinh và kiểm tra sĩ số.", "M06-01-xep-lop-phan-cong.png"),
        ("B9", "Người thực hiện: Học vụ. Lập và kiểm tra thời khóa biểu từ ngày hiện tại.", "M07-01-thoi-khoa-bieu.png"),
        ("B10", "Người thực hiện: Giáo viên/Học vụ. Mở buổi học và thực hiện điểm danh.", "M07-02-diem-danh.png"),
        ("B11", "Người thực hiện: Học vụ/Giáo viên. Tiếp nhận và xử lý đơn xin phép.", "M08-01-don-xin-phep.png"),
        ("B12", "Người thực hiện: Giáo viên/Học vụ. Ghi nhận trao đổi với phụ huynh.", "M08-02-trao-doi-phu-huynh.png"),
        ("B13", "Người thực hiện: Quản lý/Học vụ. Tạo thông báo và chọn đúng phạm vi nhận.", "M09-01-thong-bao.png"),
        ("B14", "Người thực hiện: Kế toán. Tạo kỳ thu, khoản thu và theo dõi công nợ.", "M10-01-ky-thu-cong-no.png"),
        ("B15", "Người thực hiện: Kế toán; Quản lý duyệt. Theo dõi hoàn phí, chuyển phí hoặc bảo lưu.", "M10-02-dieu-chinh.png"),
        ("B16", "Người thực hiện: Kế toán đề xuất; Quản lý duyệt. Ghi nhận chi phí vận hành.", "M10-03-chi-phi.png"),
        ("B17", "Người thực hiện: Giáo viên/Học vụ. Mở hồ sơ để cập nhật trạng thái, học tập và thành tích.", "M11-01-ket-qua-hoc-tap.png"),
        ("B18", "Người thực hiện: Quản lý/Kế toán. Đối chiếu báo cáo tài chính và số liệu chi tiết.", "M12-01-bao-cao-tai-chinh.png"),
    ]
    for step_code, description, image_name in ui_steps:
        doc.add_heading(f"{step_code} — {description}", level=2)
        add_figure(doc, image_name, f"Hình {step_code}. {description}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
