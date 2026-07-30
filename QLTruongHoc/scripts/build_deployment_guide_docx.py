from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "HUONG_DAN_TRIEN_KHAI_VA_PHAT_HANH_QLTRUONGHOC_CPANEL.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF4CE"
RISK = "FDECEC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if table.rows:
        set_repeat_table_header(table.rows[0])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=11, bold=False, color=INK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_para(doc, text="", bold_prefix=None, style=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for i, line in enumerate(code.strip().splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9, color="202124", name="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_run_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Hạng mục"
    hdr[1].text = "Giá trị / hướng dẫn"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for idx, cell in enumerate(cells):
            for run in cell.paragraphs[0].runs:
                set_run_font(run, bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run("☐ " + item)
        set_run_font(r)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def new_page(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.82)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.4)
section.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(4)
    styles[name].paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.text = "QLTRUONGHOC | HƯỚNG DẪN TRIỂN KHAI VÀ PHÁT HÀNH"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    set_run_font(run, 8.5, bold=True, color=MUTED)
footer = section.footer
add_page_number(footer.paragraphs[0])

# Cover
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HƯỚNG DẪN TRIỂN KHAI\nVÀ PHÁT HÀNH PHẦN MỀM")
set_run_font(r, 27, bold=True, color=INK)
p.paragraph_format.space_after = Pt(10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HỆ THỐNG QUẢN LÝ TRƯỜNG HỌC QLTRUONGHOC")
set_run_font(r, 16, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(28)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Triển khai tại vireon.vn trên cPanel / CloudLinux / LiteSpeed Passenger")
set_run_font(r, 12.5, color=MUTED)
p.paragraph_format.space_after = Pt(72)
add_kv_table(doc, [
    ("Đơn vị sở hữu", "Paul Digital / Vireon"),
    ("Tác giả", "Nhóm dự án QLTruongHoc"),
    ("Phiên bản tài liệu", "1.0"),
    ("Ngày phát hành", "30/07/2026"),
    ("Phạm vi", "Triển khai lần đầu, vận hành, hotfix, bugfix, CR và rollback"),
])

new_page(doc)
heading(doc, "Mục lục", 1)
for text in [
    "1. Mục đích, phạm vi và kiến trúc",
    "2. Chuẩn bị trước khi triển khai",
    "3. Build và tạo deploy package",
    "4. Cấu hình domain, Node.js App và thư mục",
    "5. Tạo và nhập cơ sở dữ liệu",
    "6. Cấu hình biến môi trường và dependency",
    "7. Khởi động, nghiệm thu và kiểm thử",
    "8. Xử lý sự cố thường gặp",
    "9. Quy trình phát hành hotfix, bugfix và CR",
    "10. Quy trình rollback",
    "11. Backup, bảo mật và vận hành định kỳ",
    "12. Checklist triển khai nhanh",
]:
    add_bullet(doc, text)

heading(doc, "Quy ước sử dụng", 2)
add_para(doc, "Mỗi bước gồm: người thực hiện, thao tác, kết quả mong đợi và kịch bản verify. Chỉ chuyển sang bước kế tiếp khi verify đạt.")
add_callout(doc, "Nguyên tắc an toàn", "Không đưa source, private key, .env.local hoặc mật khẩu vào deploy package. Không chạy script reset/full_reset trên production. Không dùng npm audit fix --force.", CAUTION)

heading(doc, "1. Mục đích, phạm vi và kiến trúc", 1)
heading(doc, "1.1. Mục đích", 2)
add_para(doc, "Tài liệu hướng dẫn triển khai QLTruongHoc lên hosting cPanel dùng chung cho paul-digitalhub.com và addon domain vireon.vn, đồng thời quy định cách phát hành các bản sửa lỗi và Change Request (CR) sau khi hệ thống đi vào vận hành.")
heading(doc, "1.2. Kiến trúc thực tế", 2)
add_code(doc, """
Internet
  → LiteSpeed / cPanel
     → vireon.vn/                         : website hiện tại
     → vireon.vn/wp/                      : WordPress
     → vireon.vn/app-portal/edu-management/
        → CloudLinux Node.js Selector / Passenger
        → /home/pauldigi/apps/edu-management
        → MariaDB: pauldigi_edu_management
""")
add_callout(doc, "Điểm quan trọng", "Ứng dụng Node.js đặt ngoài document root /home/pauldigi/vireon.vn. URL công khai do Passenger ánh xạ tới Application Root; không tạo app bên trong thư mục WordPress.", LIGHT_BLUE)

heading(doc, "1.3. Vai trò thực hiện", 2)
add_kv_table(doc, [
    ("Lập trình viên", "Sửa code, kiểm thử, build, tạo package, cung cấp release note."),
    ("Quản trị hosting", "Tạo Node.js App, database, environment variables, upload và restart."),
    ("Người kiểm thử", "Thực hiện smoke test và regression theo checklist."),
    ("Chủ hệ thống", "Phê duyệt CR, thời gian triển khai, rollback và nghiệm thu."),
])

heading(doc, "2. Chuẩn bị trước khi triển khai", 1)
heading(doc, "2.1. Kiểm tra domain và document root", 2)
add_para(doc, "Người thực hiện: Quản trị hosting.")
add_code(doc, """
whoami
pwd
uapi --output=jsonpretty DomainInfo single_domain_data domain=vireon.vn
ls -la /home/pauldigi/vireon.vn
""")
add_para(doc, "Kết quả mong đợi: user pauldigi; home /home/pauldigi; documentroot /home/pauldigi/vireon.vn; thư mục wp tồn tại.")
heading(doc, "2.2. Kiểm tra Node.js và tính năng cPanel", 2)
add_code(doc, """
/opt/alt/alt-nodejs22/root/usr/bin/node --version
/opt/alt/alt-nodejs22/root/usr/bin/npm --version
uapi --output=jsonpretty Features list_features | grep -i -C 3 node
""")
add_para(doc, "Kết quả mong đợi: Node.js 22.x, npm hoạt động và cPanel có Setup Node.js App.")
heading(doc, "2.3. Điều kiện trước khi phát hành", 2)
add_checklist(doc, [
    "Đã xác định thời gian bảo trì và người chịu trách nhiệm.",
    "Đã backup database và thư mục uploads.",
    "Đã ghi nhận package/version đang chạy để rollback.",
    "Đã có mật khẩu database và quyền cPanel nhưng không ghi vào tài liệu hoặc ticket công khai.",
    "Đã kiểm tra dung lượng hosting và giới hạn tài nguyên CloudLinux.",
])

heading(doc, "3. Build và tạo deploy package", 1)
heading(doc, "3.1. Thành phần package", 2)
add_kv_table(doc, [
    ("app_wrapper.cjs", "Startup file CommonJS để Passenger nạp backend ESM."),
    ("dist-client/", "Frontend đã build với base path /app-portal/edu-management/."),
    ("dist-server/", "Backend JavaScript đã biên dịch."),
    ("package.json", "Chỉ chứa dependency production."),
    ("package-lock.json", "Khóa phiên bản dependency cho npm."),
    ("uploads/", "Thư mục giữ chỗ; không ghi đè/xóa dữ liệu upload khi cập nhật."),
    ("deploy/", "Mẫu môi trường và hướng dẫn cPanel."),
])
heading(doc, "3.2. Lệnh kiểm tra và tạo package", 2)
add_para(doc, "Người thực hiện: Lập trình viên, trên máy build.")
add_code(doc, """
pnpm install --frozen-lockfile
pnpm typecheck
pnpm test
$env:NODE_ENV='production'
$env:VITE_APP_BASE_PATH='/app-portal/edu-management/'
pnpm build
./scripts/create-cpanel-package.ps1
""")
add_para(doc, "Kết quả mong đợi: typecheck đạt, toàn bộ test đạt, build thành công và có ZIP qltruonghoc-cpanel-vireon-YYYYMMDD-HHMMSS.zip.")
heading(doc, "3.3. Những file không đưa lên hosting", 2)
add_code(doc, """
client/src
server/*.ts
tests
docs nội bộ
.git
node_modules từ Windows
.env.local
SSH private key
""")

heading(doc, "4. Cấu hình domain, Node.js App và thư mục", 1)
heading(doc, "4.1. Tạo Application Root", 2)
add_code(doc, """
mkdir -p /home/pauldigi/apps/edu-management
ls -ld /home/pauldigi/apps/edu-management
""")
add_para(doc, "Verify: owner/group là pauldigi; không đặt ứng dụng trong /home/pauldigi/vireon.vn.")
heading(doc, "4.2. Upload và giải nén package", 2)
add_code(doc, """
scp -i "D:\\Working\\Hosting\\SSH-key\\ed25519" PACKAGE.zip \
  pauldigi@103.200.23.126:/home/pauldigi/apps/edu-management/

cd /home/pauldigi/apps/edu-management
unzip -o PACKAGE.zip
""")
add_para(doc, "Verify:")
add_code(doc, """
test -f app_wrapper.cjs
test -f package.json
test -f dist-client/index.html
test -f dist-server/server/index.js
grep -F '/app-portal/edu-management/assets/' dist-client/index.html
""")
heading(doc, "4.3. Tạo Node.js App", 2)
add_kv_table(doc, [
    ("Node.js version", "22.23.0 hoặc bản 22.x được hosting hỗ trợ"),
    ("Application mode", "Production"),
    ("Application root", "apps/edu-management"),
    ("Domain", "vireon.vn"),
    ("Application URI", "app-portal/edu-management"),
    ("Startup file", "app_wrapper.cjs"),
])
add_para(doc, "Thao tác: cPanel → Software → Setup Node.js App → Create Application → nhập đúng bảng trên → Create.")
add_callout(doc, "Không dùng PM2", "CloudLinux Node.js Selector sử dụng Passenger. Không cần và không nên chạy PM2/Nginx riêng trong shared hosting.", CAUTION)

heading(doc, "5. Tạo và nhập cơ sở dữ liệu", 1)
heading(doc, "5.1. Tạo database và user", 2)
add_para(doc, "Vào cPanel → Database Wizard.")
add_kv_table(doc, [
    ("Database", "pauldigi_edu_management"),
    ("Database user", "pauldigi_edu_app"),
    ("Host", "localhost"),
    ("Quyền", "ALL PRIVILEGES trên đúng database"),
])
add_para(doc, "Verify bằng SSH:")
add_code(doc, """
mysql -h localhost -u pauldigi_edu_app -p pauldigi_edu_management
SELECT DATABASE(), VERSION();
exit
""")
heading(doc, "5.2. Import baseline/demo data", 2)
add_para(doc, "Vào phpMyAdmin → chọn pauldigi_edu_management → Import → chọn file qltruonghoc-demo-database-YYYYMMDD.sql → charset utf-8 → Import.")
add_para(doc, "Verify:")
add_code(doc, """
SELECT COUNT(*) AS soBang
FROM information_schema.tables
WHERE table_schema = 'pauldigi_edu_management'
  AND table_type = 'BASE TABLE';
""")
add_para(doc, "Kết quả mong đợi: 44 bảng.")
heading(doc, "5.3. Chuẩn hóa chữ hoa/thường trên Linux", 2)
add_para(doc, "Nếu dump tạo bảng chữ thường nhưng code dùng tên CamelCase, import deploy/cpanel/fix-linux-table-case.sql. Sau đó kiểm tra:")
add_code(doc, """
SELECT TABLE_NAME
FROM information_schema.tables
WHERE TABLE_SCHEMA = 'pauldigi_edu_management'
  AND BINARY TABLE_NAME IN ('NguoiDung','DonVi','VaiTro','HocSinh','GiaoVien','LopHoc');
""")
add_para(doc, "Kết quả mong đợi: đủ 6 tên bảng đúng kiểu chữ.")
add_callout(doc, "Cảnh báo", "Chỉ import script đổi tên bảng sau bản dump từ MySQL Windows. Không chạy reset/full_reset trên production.", RISK)

heading(doc, "6. Cấu hình biến môi trường và dependency", 1)
heading(doc, "6.1. Environment Variables trong Node.js App", 2)
add_kv_table(doc, [
    ("DATABASE_HOST", "localhost"),
    ("DATABASE_PORT", "3306"),
    ("DATABASE_USER", "pauldigi_edu_app"),
    ("DATABASE_PASSWORD", "Mật khẩu thật; không ghi vào tài liệu/repository"),
    ("DATABASE_NAME", "pauldigi_edu_management"),
    ("DATABASE_CONNECTION_LIMIT", "10"),
    ("AUTH_COOKIE_NAME", "edu_management_session"),
    ("COOKIE_PATH", "/app-portal/edu-management"),
    ("APP_BASE_PATH", "/app-portal/edu-management"),
])
add_para(doc, "Không đặt PORT/HOST thủ công trong Node.js Selector; Passenger quản lý tiến trình HTTP. Application Mode tự đặt NODE_ENV=production.")
heading(doc, "6.2. Kích hoạt virtual environment và cài dependency", 2)
add_code(doc, """
source /home/pauldigi/nodevenv/apps/edu-management/22/bin/activate
cd /home/pauldigi/apps/edu-management
npm install --omit=dev
npm list express mysql2 drizzle-orm --depth=0
npm audit --omit=dev
""")
add_para(doc, "Kết quả mong đợi: dependency hợp lệ; drizzle-orm từ 0.45.2 trở lên; audit 0 vulnerability.")
add_callout(doc, "Không dùng --force", "Nếu audit báo lỗi, xác định package và nâng có kiểm soát trên source; chạy lại typecheck/test/build. Không chạy npm audit fix --force trên production.", CAUTION)

heading(doc, "7. Khởi động, nghiệm thu và kiểm thử", 1)
heading(doc, "7.1. Khởi động", 2)
add_para(doc, "Trong Setup Node.js App: START APP → chờ 10–20 giây → RESTART.")
heading(doc, "7.2. Kiểm tra API và asset", 2)
add_code(doc, """
curl -i https://vireon.vn/app-portal/edu-management/api/health
curl -I https://vireon.vn/app-portal/edu-management/assets/TEN_FILE_JS.js
""")
add_para(doc, "Kết quả mong đợi: API HTTP/2 200, databaseName đúng; JS HTTP/2 200 và content-type text/javascript hoặc application/javascript.")
heading(doc, "7.3. Smoke test giao diện", 2)
add_checklist(doc, [
    "Mở trang bằng Incognito; trang đăng nhập có đầy đủ CSS.",
    "Đăng nhập admin và đổi mật khẩu tạm ngay.",
    "Cookie có Secure, HttpOnly, SameSite=Lax, Path=/app-portal/edu-management.",
    "Refresh trang con không 404 và không mất phiên.",
    "Mở trang con trong tab mới vẫn giữ đúng base path.",
    "Upload/xem/tải file hoạt động.",
    "Phiếu in hiển thị logo và dữ liệu.",
    "vireon.vn/ và vireon.vn/wp/ không bị ảnh hưởng.",
])
heading(doc, "7.4. Regression nghiệp vụ tối thiểu", 2)
add_checklist(doc, [
    "Quản trị hệ thống: dashboard, đơn vị, tài khoản, phân quyền.",
    "Quản lý đơn vị: nhân sự và phạm vi dữ liệu đúng đơn vị.",
    "Tuyển sinh: lead, liên hệ, xác nhận nhập học.",
    "Học vụ: chương trình, lớp, xếp lớp, lịch học.",
    "Kế toán: kỳ thu, khoản phải thu, phiếu thu, phiếu in.",
    "Giáo viên: lịch, điểm danh, báo giảng, đánh giá.",
    "Phụ huynh: hai con ở hai đơn vị nhận đủ thông báo, xem trao đổi và xin phép.",
])

heading(doc, "8. Xử lý sự cố thường gặp", 1)
add_kv_table(doc, [
    ("Trang trắng, MIME text/html", "Passenger giữ nguyên Application URI. Đặt APP_BASE_PATH và dùng backend có middleware chuẩn hóa base path; restart app."),
    ("API 500/503", "Kiểm tra Environment Variables, trạng thái Node.js App, database và error_log."),
    ("node/npm not found", "Kích hoạt /home/pauldigi/nodevenv/apps/edu-management/22/bin/activate."),
    ("Failed query from NguoiDung", "Kiểm tra kiểu chữ tên bảng Linux; import fix-linux-table-case.sql."),
    ("Sai mật khẩu admin", "Reset hash bằng quy trình được phê duyệt; mở khóa; bắt buộc đổi mật khẩu."),
    ("npm audit thiếu lockfile", "Đưa package-lock vào release hoặc tạo npm install --package-lock-only; không dùng --force."),
])
heading(doc, "8.1. Thu thập bằng chứng lỗi", 2)
add_code(doc, """
curl -i https://vireon.vn/app-portal/edu-management/api/health
tail -n 100 /home/pauldigi/error_log
find /home/pauldigi -maxdepth 3 -type f -iname '*passenger*.log' -print
""")
add_para(doc, "Khi báo lỗi cần ghi: thời điểm, URL, tài khoản/role (không ghi mật khẩu), request status, ảnh Network/Console, bước tái hiện và phiên bản release.")

heading(doc, "9. Quy trình phát hành hotfix, bugfix và CR", 1)
heading(doc, "9.1. Phân loại phát hành", 2)
add_kv_table(doc, [
    ("Hotfix", "Lỗi production nghiêm trọng; phạm vi nhỏ; ưu tiên khôi phục dịch vụ."),
    ("Bugfix release", "Sửa một hoặc nhiều lỗi đã xác nhận; có regression test."),
    ("CR nhỏ", "Thay đổi chức năng ít ảnh hưởng dữ liệu/quyền; phát hành theo lịch."),
    ("CR lớn", "Thay đổi quy trình, schema, quyền hoặc nhiều module; bắt buộc staging/UAT."),
])
heading(doc, "9.2. Chu kỳ CR chuẩn", 2)
for text in [
    "Tiếp nhận: ghi mục tiêu, người dùng, hiện trạng, kết quả mong muốn và tiêu chí nghiệm thu.",
    "Phân tích ảnh hưởng: frontend, API, database, phân quyền, dữ liệu hai loại đơn vị, tài liệu.",
    "Phê duyệt: xác nhận phạm vi, độ ưu tiên, thời gian và kế hoạch rollback.",
    "Phát triển: tạo branch, code, migration tiến tới (forward-only), test.",
    "Kiểm thử: typecheck, automated test, module test, role/unit regression, UAT.",
    "Đóng gói: build production, tạo ZIP, checksum, release note.",
    "Triển khai: backup, stop app, upload, migration, npm install, start/restart.",
    "Nghiệm thu: smoke test, business test, theo dõi log, xác nhận hoàn tất.",
]:
    add_bullet(doc, text)
heading(doc, "9.3. Quy tắc version", 2)
add_para(doc, "Khuyến nghị Semantic Versioning MAJOR.MINOR.PATCH:")
add_bullet(doc, "PATCH: hotfix/bugfix không thay đổi tương thích, ví dụ 0.4.1.")
add_bullet(doc, "MINOR: CR bổ sung chức năng tương thích, ví dụ 0.5.0.")
add_bullet(doc, "MAJOR: thay đổi lớn hoặc không tương thích, ví dụ 1.0.0.")
heading(doc, "9.4. Release note bắt buộc", 2)
add_checklist(doc, [
    "Mã release, ngày giờ và người triển khai.",
    "Danh sách ticket/CR và mô tả thay đổi.",
    "Danh sách file/module bị ảnh hưởng.",
    "Migration và yêu cầu cấu hình mới.",
    "Test đã chạy và kết quả.",
    "Hướng dẫn deploy và verify.",
    "Rủi ro đã biết và rollback.",
])
heading(doc, "9.5. Các bước deploy bản sửa/CR", 2)
add_code(doc, """
# Trên máy build
pnpm install --frozen-lockfile
pnpm typecheck
pnpm test
./scripts/create-cpanel-package.ps1

# Trên hosting
STOP APP
backup database + uploads
upload ZIP mới
cd /home/pauldigi/apps/edu-management
unzip -o PACKAGE_MOI.zip
source /home/pauldigi/nodevenv/apps/edu-management/22/bin/activate
npm install --omit=dev
START APP → RESTART
""")
add_para(doc, "Nếu release có migration: chạy migration sau backup và trước START APP. Migration phải có mã riêng, idempotent khi có thể, không DROP dữ liệu và đã test trên bản sao.")
heading(doc, "9.6. Verify sau release", 2)
add_checklist(doc, [
    "Health API 200 và kết nối đúng database.",
    "Asset JS đúng MIME.",
    "Đăng nhập, cookie và refresh route hoạt động.",
    "Test đúng ticket/CR.",
    "Regression các module phụ thuộc.",
    "Log không có lỗi mới trong ít nhất 15–30 phút.",
])

heading(doc, "10. Quy trình rollback", 1)
heading(doc, "10.1. Khi nào rollback", 2)
add_bullet(doc, "Health API hoặc đăng nhập không hoạt động sau thời gian xử lý cho phép.")
add_bullet(doc, "Lỗi dữ liệu, phân quyền hoặc ảnh hưởng website vireon.vn/wp.")
add_bullet(doc, "CR không đạt tiêu chí nghiệm thu hoặc có lỗi nghiêm trọng.")
heading(doc, "10.2. Rollback code", 2)
add_code(doc, """
STOP APP
khôi phục dist-client, dist-server, package.json, package-lock.json,
app_wrapper.cjs từ package release trước
npm install --omit=dev
START APP → RESTART
""")
heading(doc, "10.3. Rollback database", 2)
add_para(doc, "Chỉ rollback database khi migration đã thay đổi dữ liệu/schema và có bản backup ngay trước release. Thực hiện trong cửa sổ bảo trì, xác nhận đúng database, ghi log thời điểm và người phê duyệt.")
add_callout(doc, "Không rollback mù", "Nếu hệ thống đã phát sinh dữ liệu mới sau release, restore toàn bộ database có thể làm mất dữ liệu. Cần đánh giá chênh lệch và chọn script khôi phục phù hợp.", RISK)
heading(doc, "10.4. Verify rollback", 2)
add_checklist(doc, [
    "Health API 200.",
    "Đăng nhập được bằng tài khoản hợp lệ.",
    "Các nghiệp vụ cốt lõi hoạt động.",
    "Phiên bản giao diện/asset đúng release cũ.",
    "Không có lỗi database mới.",
])

heading(doc, "11. Backup, bảo mật và vận hành định kỳ", 1)
heading(doc, "11.1. Backup", 2)
add_checklist(doc, [
    "Database hằng ngày; giữ nhiều phiên bản theo chính sách.",
    "Uploads hằng ngày hoặc theo tần suất phát sinh.",
    "Deploy package của từng release.",
    "Export Environment Variables an toàn hoặc lưu trong password manager.",
    "Thử restore định kỳ, không chỉ kiểm tra file backup tồn tại.",
])
heading(doc, "11.2. Bảo mật", 2)
add_checklist(doc, [
    "Đổi toàn bộ mật khẩu demo trước khi public.",
    "Không dùng chung MySQL user với WordPress.",
    "Cookie Secure, HttpOnly, SameSite=Lax và đúng Path.",
    "Không commit .env.local/private key.",
    "Theo dõi npm audit nhưng nâng dependency có kiểm thử.",
    "Khóa hoặc xóa tài khoản không còn sử dụng.",
])
heading(doc, "11.3. Giám sát", 2)
add_checklist(doc, [
    "Theo dõi health API.",
    "Theo dõi error_log/Passenger log.",
    "Theo dõi CPU, RAM, process, inode và dung lượng CloudLinux.",
    "Theo dõi dung lượng database và uploads.",
    "Kiểm tra chứng chỉ SSL và ngày hết hạn.",
])

heading(doc, "12. Checklist triển khai nhanh", 1)
add_checklist(doc, [
    "Package có app_wrapper.cjs, dist-client, dist-server, package.json và package-lock.json.",
    "Đã backup database/uploads và giữ package cũ.",
    "Node.js App đang STOP trước khi thay code.",
    "Đã giải nén đúng /home/pauldigi/apps/edu-management.",
    "Environment Variables đầy đủ, gồm APP_BASE_PATH.",
    "npm install --omit=dev thành công; audit đạt.",
    "Migration/baseline đã import đúng và tên bảng đúng kiểu chữ.",
    "START APP và RESTART thành công.",
    "Health API 200; asset đúng MIME.",
    "Đăng nhập, cookie, refresh route và upload hoạt động.",
    "Website vireon.vn/ và /wp/ không bị ảnh hưởng.",
    "Release note, kết quả test và người nghiệm thu đã được ghi nhận.",
])

heading(doc, "Phụ lục A — Mẫu biên bản phát hành", 1)
add_kv_table(doc, [
    ("Release", "........................................................"),
    ("Loại", "Hotfix / Bugfix / CR nhỏ / CR lớn"),
    ("Thời gian", "........................................................"),
    ("Người triển khai", "........................................................"),
    ("Ticket/CR", "........................................................"),
    ("Package", "........................................................"),
    ("Database migration", "Có / Không; mã: ................................"),
    ("Backup", "Database: ............ | Uploads: ............"),
    ("Kết quả smoke test", "Đạt / Không đạt"),
    ("Kết quả nghiệp vụ", "Đạt / Không đạt"),
    ("Rollback", "Không cần / Đã rollback lúc ............"),
    ("Người nghiệm thu", "........................................................"),
])

heading(doc, "Phụ lục B — Mẫu ghi nhận lỗi sau phát hành", 1)
add_kv_table(doc, [
    ("Thời điểm", "........................................................"),
    ("URL/module", "........................................................"),
    ("Role/đơn vị", "........................................................"),
    ("Các bước tái hiện", "........................................................"),
    ("Kết quả thực tế", "........................................................"),
    ("Kết quả mong đợi", "........................................................"),
    ("HTTP/Console/Log", "........................................................"),
    ("Mức độ", "Blocker / High / Medium / Low"),
    ("Quyết định", "Fix tiếp / Rollback / Theo dõi"),
])

doc.core_properties.title = "Hướng dẫn triển khai và phát hành QLTruongHoc trên cPanel"
doc.core_properties.subject = "SOP triển khai, release, CR và rollback"
doc.core_properties.author = "Nhóm dự án QLTruongHoc - Paul Digital / Vireon"
doc.core_properties.keywords = "QLTruongHoc, cPanel, CloudLinux, Passenger, deployment, release, CR"
doc.core_properties.comments = "Phiên bản 1.0 - 30/07/2026"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
